"""Rebuild references.docx from the corrected reference list."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
import copy

REFS = [
    # Format: display text (hanging indent applied programmatically)
    ("Bergé, Laurent (2022). ",
     "fixest",
     ": Fast Fixed Effects Estimation. R package version 0.10.4. https://lrberge.github.io/fixest/."),

    ("Callaway, Brantly, and Pedro H. C. Sant'Anna (2021). "
     "\"Difference-in-Differences with Multiple Time Periods.\" "
     "Journal of Econometrics 225(2): 200–230. "
     "https://doi.org/10.1016/j.jeconom.2020.12.001.",
     None, None),

    ("Currie, Janet, and Soodeh Saberian (2025). "
     "\"Wildfire, Smoke and Mental Health in Canada.\" "
     "NBER Working Paper No. 33912. "
     "https://www.nber.org/papers/w33912.",
     None, None),

    ("Du, Ran, Ke Liu, Dangru Zhao, and Qiyun Fang (2024). "
     "\"Climate Disaster and Cognitive Ability: Evidence From Wildfire.\" "
     "International Journal of Public Health 69: 1607128. "
     "https://doi.org/10.3389/ijph.2024.1607128.",
     None, None),

    ("Greenberg, Paul E., Andree-Anne Fournier, Tammy Sisitsky, "
     "Mark Simes, Richard Berman, Sarah H. Koenigsberg, and Ronald C. Kessler (2021). "
     "\"The Economic Burden of Adults with Major Depressive Disorder in the United States "
     "(2010 and 2018).\" PharmacoEconomics 39(6): 653–665. "
     "https://doi.org/10.1007/s40273-021-01019-4.",
     None, None),

    ("Imbens, Guido W., and Michal Kolesár (2016). "
     "\"Robust Standard Errors in Small Samples: Some Practical Advice.\" "
     "Review of Economics and Statistics 98(4): 701–712. "
     "https://doi.org/10.1162/REST_a_00552.",
     None, None),

    ("Jung, Youn Soo, Mary M. Johnson, Marshall Burke, Sara Heft-Neal, "
     "Mathieu Bondy, Sayantani Chinthrajah, Ashley C. Cullen, Kim Nelson, "
     "Nina Dresser, and Kari Nadeau (2025). "
     "\"Fine Particulate Matter From 2020 California Wildfires and Mental "
     "Health–Related Emergency Department Visits.\" "
     "JAMA Network Open 8(4): e253326. "
     "https://doi.org/10.1001/jamanetworkopen.2025.3326.",
     None, None),

    ("Merdjanoff, Alexis A., Kate Burrows, and Kathleen Lynch (2026). "
     "\"Projecting the Mental Health Impacts of the 2025 LA Wildfires: "
     "Lessons from Hurricane Katrina.\" "
     "Environmental Research Letters 21(11): 111004. "
     "https://doi.org/10.1088/1748-9326/ae6d16.",
     None, None),

    ("Pustejovsky, James E. (2022). ",
     "clubSandwich",
     ": Cluster-Robust (Sandwich) Variance Estimators with Small-Sample "
     "Corrections. R package version 0.5.8. "
     "https://CRAN.R-project.org/package=clubSandwich."),

    ("Wettstein, Zachary S., and Ambarish Vaidyanathan (2024). "
     "\"Psychotropic Medication Prescriptions and Large California Wildfires.\" "
     "JAMA Network Open 7(2): e2356466. "
     "https://doi.org/10.1001/jamanetworkopen.2023.56466.",
     None, None),

    ("Ye, Xinyue, Yuning Ye, Xiao Huang, and Tracy Onega (2026). "
     "\"Wildfires and Public Health: A Comprehensive Review of Human-Centric Studies.\" "
     "GeoHealth 10(2): e2025GH001534. "
     "https://doi.org/10.1029/2025GH001534.",
     None, None),
]

doc = Document()

# ── Page margins matching main paper ──────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ── References heading ─────────────────────────────────────────────────────────
h = doc.add_paragraph("References")
h.style = doc.styles["Heading 1"]
run = h.runs[0]
run.font.size = Pt(12)
run.font.bold = True
run.font.name = "Times New Roman"

# ── Helper: paragraph format for hanging indent ────────────────────────────────
def set_hanging(para, size_pt=12):
    pf = para.paragraph_format
    pf.left_indent        = Inches(0.5)
    pf.first_line_indent  = Inches(-0.5)
    pf.space_after        = Pt(0)
    pf.space_before       = Pt(6)
    pf.line_spacing_rule  = WD_LINE_SPACING.DOUBLE

def add_run(para, text, italic=False, size_pt=12):
    r = para.add_run(text)
    r.font.size = Pt(size_pt)
    r.font.name = "Times New Roman"
    r.font.italic = italic
    return r

# ── Add each reference ─────────────────────────────────────────────────────────
for entry in REFS:
    before, italic_part, after = entry
    p = doc.add_paragraph()
    set_hanging(p)

    if italic_part is None:
        # Plain text entry
        add_run(p, before)
    else:
        # Entry with an italic software name
        add_run(p, before)
        add_run(p, italic_part, italic=True)
        add_run(p, after)

out = r"C:\Users\chenyon\Research Paper 2026(1)\references.docx"
doc.save(out)
print(f"Saved: {out}")
