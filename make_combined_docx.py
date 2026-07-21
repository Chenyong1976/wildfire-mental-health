"""
Assembles the complete paper from individual section .docx files.
Uses XML element copying to preserve all formatting.
"""
import copy
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PAPER_DIR = r"C:\Users\chenyon\Research Paper 2026(1)"

# ── Section files in order ────────────────────────────────────────────────────
SECTIONS = [
    "introduction.docx",
    "section_data.docx",
    "section_empirical.docx",
    "section_results.docx",
    "section_robustness.docx",
    "section_conclusion.docx",
]

# ── Build master from the introduction (it has the title block) ───────────────
master = Document(PAPER_DIR + "\\" + SECTIONS[0])

# Insert abstract after the title block (before the Introduction section heading)
# Find the first non-title paragraph index
def insert_abstract(doc):
    """Insert an abstract block after the title block paragraphs."""
    # Locate the first section heading ("1.  Introduction" or similar)
    insert_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("1.") or para.text.strip() == "Introduction":
            insert_idx = i
            break
    if insert_idx is None:
        return  # fallback: do nothing

    abstract_text = (
        "I estimate the causal effect of the 2015 Western U.S. wildfire season on "
        "county-level mental health using a matched difference-in-differences design. "
        "Wildfire Hazard Potential matching identifies 81 treated counties from "
        "structurally similar never-burned controls within fire-risk strata. "
        "Four years post-fire, treated counties report depression prevalence 2.76 "
        "percentage points higher (SE = 0.94, p = 0.004) and poor mental health days "
        "1.63 percentage points higher (SE = 0.51, p = 0.002). A dose-response "
        "specification replacing the binary treatment indicator with log burn intensity "
        "yields a coefficient of 1.18 per log-unit (p = 0.020), inconsistent with "
        "selection confounding. Suicide and overdose mortality estimates are consistently "
        "positive but imprecise, reflecting widespread CDC WONDER cell suppression in "
        "small Western counties. Non-metropolitan counties drive the mortality signal: "
        "the four-year suicide event-study coefficient is significant among rural "
        "counties (p = 0.038), consistent with mental health provider access as the "
        "operative mechanism."
    )

    # Insert abstract label paragraph
    from docx.oxml import OxmlElement
    body = doc.element.body
    ref_elem = doc.paragraphs[insert_idx]._element

    def make_para(text, bold=False, indent=False, italic=False):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        rPr_elem = OxmlElement('w:rPr')
        if bold:
            b = OxmlElement('w:b')
            rPr_elem.append(b)
        pPr_spacing = OxmlElement('w:spacing')
        pPr_spacing.set(qn('w:line'), '480')
        pPr_spacing.set(qn('w:lineRule'), 'auto')
        pPr_spacing.set(qn('w:after'), '0')
        pPr.append(pPr_spacing)
        if indent:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:firstLine'), '720')
            pPr.append(ind)
        p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')
        rPr.append(sz)
        if bold:
            b2 = OxmlElement('w:b')
            rPr.append(b2)
        if italic:
            i_elem = OxmlElement('w:i')
            rPr.append(i_elem)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        return p

    # Insert: blank, Abstract label, abstract text, JEL line, blank
    elems = [
        make_para(""),
        make_para("Abstract", bold=True),
        make_para(abstract_text),
        make_para("JEL codes: I12, I18, Q54, R11", italic=True),
        make_para("Keywords: wildfire, mental health, depression, difference-in-differences, rural health, natural disaster", italic=True),
        make_para(""),
    ]
    for elem in reversed(elems):
        ref_elem.addprevious(elem)


insert_abstract(master)

# ── Append remaining sections ─────────────────────────────────────────────────
def append_section(master_doc, section_path):
    """Copy all body elements from source doc into master, skipping its title block."""
    src = Document(section_path)
    src_body = src.element.body

    # Detect how many paragraphs constitute the title block in section files
    # (title line + blank line + institution line + blank line = 4 elements)
    # We identify them by bold+centered text containing the paper title keyword
    skip_until = 0
    for i, child in enumerate(src_body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            text = ''.join(t.text or '' for t in child.iter(qn('w:t')))
            # Stop skipping once we hit a section heading (bold, not centered title)
            if text.strip() and 'Wildfire and Long-Run' not in text and '[Author' not in text:
                skip_until = i
                break
        # Also skip sectPr (section properties) at the very end

    master_body = master_doc.element.body
    # Insert before the final sectPr
    sect_pr = master_body.find(qn('w:sectPr'))

    for i, child in enumerate(src_body):
        if i < skip_until:
            continue
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sectPr':
            continue  # skip section properties from source
        node = copy.deepcopy(child)
        if sect_pr is not None:
            sect_pr.addprevious(node)
        else:
            master_body.append(node)


for section_file in SECTIONS[1:]:
    path = PAPER_DIR + "\\" + section_file
    print(f"Appending {section_file}...")
    append_section(master, path)

# ── Save ─────────────────────────────────────────────────────────────────────
out = PAPER_DIR + "\\paper_combined.docx"
master.save(out)
import os
size_kb = os.path.getsize(out) / 1024
print(f"\nSaved: {out}  ({size_kb:.1f} KB)")
