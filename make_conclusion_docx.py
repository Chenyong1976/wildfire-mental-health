import sys; sys.path.insert(0, r'C:\Users\chenyon\Research Paper 2026(1)')
from make_equations import add_math_runs
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(24)
style.paragraph_format.space_after  = Pt(0)
style.paragraph_format.space_before = Pt(0)


def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True


def para(text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    add_math_runs(p, text)


def blank():
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_after  = Pt(0)


# ─── Section ─────────────────────────────────────────────────────────────────
h1("7.  Conclusion")
blank()

para(
    "Four years after the 2015 fire season, counties that burned show depression "
    "prevalence nearly three percentage points above comparable counties that did "
    "not: a gap large enough to register in two independent survey measures and "
    "robust to a dose-response test that is inconsistent with simple selection into "
    "treatment. Suicide and overdose mortality point estimates are consistently "
    "positive across every threshold and specification, but confidence intervals are "
    "too wide for definitive conclusions; suppression of roughly half of all "
    "county-year mortality cells is the binding constraint, not the design. The "
    "findings narrow the question from 'do wildfires harm mental health?' to 'how "
    "large is the harm and how long does it persist?' and they locate the harm "
    "specifically in non-metropolitan counties, where a four-year suicide "
    "coefficient of 0.241 IHS units achieves statistical significance that the "
    "full-sample estimate does not."
)

para(
    "The treatment effect heterogeneity by rurality has direct policy content. "
    "Rural HPSA-designated counties are simultaneously the counties that face the "
    "highest wildfire risk, the counties where the depression effect is "
    "concentrated, and the counties where behavioral health providers are most "
    "scarce. That combination carries a direct cost. If the 2.76 "
    "percentage point increase in depression prevalence among the 79 T1 treated "
    "counties represents a persistent burden (even a partial and temporary "
    "one), it translates to roughly 3,500 additional adults with untreated or "
    "undertreated depression in the affected counties at any given post-fire year, "
    "using the 2019 cross-section as a benchmark. At conservative annual cost "
    "estimates for untreated depression ($3,000-$10,000 per person in lost "
    "productivity and healthcare utilization, consistent with Greenberg et al. "
    "(2021)), the aggregate burden from a single moderate fire season across 79 "
    "counties falls in the $10-$35 million range annually, well above the "
    "marginal cost of deploying mobile crisis teams or telemedicine-based "
    "behavioral health services in the affected areas. FEMA's Crisis Counseling "
    "Assistance and Training Program and CMS's pandemic-era mental health "
    "telemedicine flexibilities both provide administrative vehicles; the policy "
    "gap is sustained funding beyond the immediate disaster window, when "
    "estimates suggest the depression burden is still growing."
)

para(
    "Two extensions would most sharpen the evidence. First, the depression finding "
    "is pinned to a single 2019 cross-section. CDC PLACES now extends through 2023, "
    "and a four-year panel of depression prevalence would enable a direct DiD for "
    "the headline outcome, replacing the cross-sectional approximation with a "
    "within-county comparison. Whether the 2.76 percentage point gap persists or "
    "decays after 2019 is the most policy-relevant open question: persistence "
    "implies a structural shift in the mental health environment of "
    "wildfire-affected communities; decay implies that the acute treatment window "
    "matters more than long-run provider supply. Second, the rurality heterogeneity "
    "is consistent with a provider-access mechanism but does not establish it "
    "causally. A follow-up study interacting the treatment with pre-fire HPSA "
    "designation and post-fire Medicaid telemedicine adoption (both of which vary "
    "across counties and over time) would test whether closing the provider gap "
    "attenuates the mental health burden. That variation exists in administrative "
    "data and could support a within-treated-county IV strategy using CMS waiver "
    "adoption as an instrument for effective provider access."
)

out = r"C:\Users\chenyon\Research Paper 2026(1)\section_conclusion.docx"
doc.save(out)
print("Saved: " + out)
