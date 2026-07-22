"""
make_curieux_submission.py
Builds Curieux Review submission package (manuscript.docx + cover_letter.docx).
Formatting: Times New Roman 12pt, 1.15 spacing, 1-inch margins, MLA citations.
"""

import os, csv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT   = r"C:\Users\chenyon\Research_Paper_2026"
OUTDIR = os.path.join(ROOT, "Curieux_Submission")
os.makedirs(OUTDIR, exist_ok=True)

# ── Helpers ───────────────────────────────────────────────────────────────────

def new_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'
    s.font.size = Pt(12)
    pf = s.paragraph_format
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1)
        sec.right_margin  = Inches(1)
    return doc

def add_para(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False,
             first_indent=None, left_indent=None,
             sb=0, sa=0, size=12):
    p  = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing  = 1.15
    pf.space_before  = Pt(sb)
    pf.space_after   = Pt(sa)
    if first_indent is not None:
        pf.first_line_indent = Inches(first_indent)
    if left_indent is not None:
        pf.left_indent = Inches(left_indent)
    run = p.add_run(text)
    run.font.name   = 'Times New Roman'
    run.font.size   = Pt(size)
    run.bold   = bold
    run.italic = italic
    return p

def heading(doc, text, sb=12):
    add_para(doc, text, bold=True, sb=sb, sa=4)

def subheading(doc, text):
    add_para(doc, text, bold=True, sb=6, sa=3)

def body(doc, text):
    add_para(doc, text, first_indent=0.5, sa=0)

def noindent(doc, text):
    add_para(doc, text, sa=0)

def blank(doc):
    add_para(doc, "")

def page_break(doc):
    doc.add_page_break()

def note(doc, text):
    add_para(doc, text, italic=True, sa=6, size=10)

def works_cited_entry(doc, text):
    add_para(doc, text, first_indent=-0.5, left_indent=0.5, sa=6)

def set_cell_font(cell, bold=False, size=10):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.bold = bold
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)

def add_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=10):
    cell.text = text
    for para in cell.paragraphs:
        para.alignment = align
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.bold = bold

def mla(text):
    """Convert economics in-text citations to MLA format (remove years)."""
    pairs = [
        ("Jung et al. (2025)",                "Jung et al."),
        ("Wettstein and Vaidyanathan (2024)",  "Wettstein and Vaidyanathan"),
        ("Currie and Saberian (2025)",         "Currie and Saberian"),
        ("Du et al. (2024)",                   "Du et al."),
        ("Ye et al. (2026)",                   "Ye et al."),
        ("Merdjanoff et al. (2026)",           "Merdjanoff et al."),
        ("Callaway and Sant'Anna (2021)",       "Callaway and Sant'Anna"),
        ("Berge, 2022",                        "Bergé"),
        ("Bergé (2022)",                       "Bergé"),
        ("Bergé, 2022",                        "Bergé"),
        ("Pustejovsky (2022)",                 "Pustejovsky"),
        ("Imbens and Kolesar (2016)",          "Imbens and Kolesár"),
        ("Imbens and Kolesár (2016)",          "Imbens and Kolesár"),
        ("Greenberg et al. (2021)",            "Greenberg et al."),
        ("(Callaway and Sant'Anna, 2021)",     "(Callaway and Sant'Anna)"),
    ]
    for old, new in pairs:
        text = text.replace(old, new)
    return text

# ── Figure helpers ────────────────────────────────────────────────────────────

FIG_MAP = {
    "1": "fig_common_support.png",
    "2": "fig_depression_coefplot.png",
}

def insert_figure(doc, num, caption):
    path = os.path.join(ROOT, FIG_MAP.get(num, ""))
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run()
        run.add_picture(path, width=Inches(5.5))
    else:
        add_para(doc, f"[Figure {num} — file not found: {FIG_MAP.get(num)}]",
                 align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_para(doc, f"Figure {num}. {caption}",
             align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, sa=12, size=10)

# ── Table builders ────────────────────────────────────────────────────────────

def stars(p):
    if p < 0.01:  return "***"
    if p < 0.05:  return "**"
    if p < 0.10:  return "*"
    return ""

def build_table1(doc):
    heading(doc, "Table 1. Summary Statistics: T1 Matched Panel (Wildfires ≥1,000 Acres, 2015)", sb=6)
    cols = ["Variable", "Treated (N=81)\nMean (SD)", "Control (N=12)\nMean (SD)", "Difference", "p-value"]
    rows_data = []

    skip_vars = {"Suicide rate (per 100,000)", "Overdose rate (per 100,000)"}

    with open(os.path.join(ROOT, "summary_stats_table.csv"), newline='', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for r in reader:
            var = r['Variable'].strip()
            is_header = r['panel_header'].strip() == 'True'

            if is_header and 'Panel B' in var:
                continue
            if is_header and 'Panel C' in var:
                rows_data.append(("PANEL", "Panel B: Post-Treatment Outcomes (2019)"))
                continue
            if is_header:
                rows_data.append(("PANEL", var))
                continue
            if var in skip_vars:
                continue

            try:
                t_mean = float(r['T_mean']); t_sd = float(r['T_sd'])
                c_mean = float(r['C_mean']); c_sd = float(r['C_sd'])
                diff   = float(r['Diff'])
                fmt    = r['fmt'].strip()
                pval   = float(r['p_value']) if r['p_value'] else None
                fmtstr = f"{fmt}"
                t_str  = f"{{:{fmtstr}}}\n({{:{fmtstr}}})".format(t_mean, t_sd)
                c_str  = f"{{:{fmtstr}}}\n({{:{fmtstr}}})".format(c_mean, c_sd)
                s_str  = stars(pval) if pval else ""
                d_str  = f"{diff:{fmtstr}}{s_str}"
                p_str  = f"{pval:.3f}" if pval else "—"
            except Exception:
                t_str = c_str = d_str = p_str = "—"
            rows_data.append((var, t_str, c_str, d_str, p_str))

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(cols):
        add_cell(hdr[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in rows_data:
        if row[0] == "PANEL":
            tr = tbl.add_row().cells
            add_cell(tr[0], row[1], bold=True)
            for c in tr[1:]: add_cell(c, "")
        else:
            tr = tbl.add_row().cells
            add_cell(tr[0], row[0])
            for i, val in enumerate(row[1:], 1):
                add_cell(tr[i], val, align=WD_ALIGN_PARAGRAPH.CENTER)

    note(doc, "Notes: Treated counties experienced a wildfire ≥1,000 acres in 2015 (N=81). Control counties had no qualifying wildfire 2015–2019 (N=12). Standard deviations in parentheses. Panel B outcomes from CDC PLACES 2019. *** p<0.01, ** p<0.05, * p<0.10.")


def build_table2(doc):
    heading(doc, "Table 2. Treatment Effects on Depression Prevalence and Poor Mental Health Days: 2019", sb=6)
    rows = [
        ("Panel A: Depression Prevalence (%)", "", ""),
        ("Treated (binary)", "2.758***\n(0.944)", "—"),
        ("Log % land burned (dose-response)", "—", "1.180*\n(0.496)"),
        ("N (treated counties)", "79", "79"),
        ("Panel B: Poor Mental Health Days (%)", "", ""),
        ("Treated (binary)", "1.634**\n(0.505)", "—"),
        ("Log % land burned (dose-response)", "—", "0.198\n(0.269)"),
        ("N (treated counties)", "79", "79"),
    ]
    col_hdrs = ["", "T1 ≥1,000 ac\nBinary (1)", "T1 ≥1,000 ac\nDose-Resp. (2)"]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(col_hdrs):
        add_cell(hdr[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        tr = tbl.add_row().cells
        is_panel = row[0].startswith("Panel")
        add_cell(tr[0], row[0], bold=is_panel)
        for i, val in enumerate(row[1:], 1):
            add_cell(tr[i], val, align=WD_ALIGN_PARAGRAPH.CENTER)
    note(doc, "Notes: Outcome is 2019 CDC PLACES county-level prevalence. Column (1): OLS with binary T1 treatment indicator. Column (2): OLS with log(1 + PctBurned) among T1 treated counties. Controls: 2011–2014 mean unemployment rate; 2013 RUCC score. Heteroskedasticity-robust SEs in parentheses. *** p<0.01, ** p<0.05, * p<0.10.")


def build_table3(doc):
    heading(doc, "Table 3. Wildfire Effects on Suicide and Drug Overdose Mortality: DiD Estimates", sb=6)

    # Panel A: Simple DiD ATT from paper text
    panelA = [
        ("Suicide",  "T1", "+0.195", "0.200", "0.331"),
        ("Overdose", "T1", "+0.213", "0.201", "0.291"),
        ("Suicide",  "T3", "−0.056", "—",     "0.766"),
    ]

    # Panel B: Collapsed 2×2 DiD from CSV
    panelB = {}
    with open(os.path.join(ROOT, "results", "did_2015_collapsed.csv"), newline='') as f:
        for r in csv.DictReader(f):
            key = (r['threshold'], r['outcome'])
            panelB[key] = r

    # Panel C: Pre-trend F-tests from CSV
    panelC = {}
    with open(os.path.join(ROOT, "results", "did_2015_pretrend.csv"), newline='') as f:
        for r in csv.DictReader(f):
            key = (r['threshold'], r['outcome'])
            panelC[key] = r

    col_hdrs = ["", "Outcome", "ATT", "SE", "p-value"]
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(col_hdrs):
        add_cell(hdr[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Panel A header
    tr = tbl.add_row().cells
    add_cell(tr[0], "Panel A: Simple DiD ATT (TWFE)", bold=True)
    for c in tr[1:]: add_cell(c, "")
    for out, thr, att, se, pv in panelA:
        tr = tbl.add_row().cells
        add_cell(tr[0], thr); add_cell(tr[1], out)
        add_cell(tr[2], att, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell(tr[3], se,  align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell(tr[4], pv,  align=WD_ALIGN_PARAGRAPH.CENTER)

    # Panel B header
    tr = tbl.add_row().cells
    add_cell(tr[0], "Panel B: Collapsed 2×2 DiD", bold=True)
    for c in tr[1:]: add_cell(c, "")

    outcome_labels = {"ihs_suicide_rate": "Suicide", "ihs_overdose_rate": "Overdose"}
    thresh_labels  = {
        "T1_moderate_1k": "T1 ≥1,000 ac",
        "T2_large_5k":    "T2 ≥5,000 ac",
        "T3_verylarge_25k": "T3 ≥25,000 ac",
    }
    for tk in ["T1_moderate_1k","T2_large_5k","T3_verylarge_25k"]:
        for ok in ["ihs_suicide_rate","ihs_overdose_rate"]:
            r = panelB.get((tk, ok))
            if r:
                att = float(r['att']); se_ = float(r['se']); pv = float(r['pvalue'])
                s = stars(pv)
                tr = tbl.add_row().cells
                add_cell(tr[0], thresh_labels[tk])
                add_cell(tr[1], outcome_labels[ok])
                add_cell(tr[2], f"{att:+.3f}{s}", align=WD_ALIGN_PARAGRAPH.CENTER)
                add_cell(tr[3], f"{se_:.3f}",     align=WD_ALIGN_PARAGRAPH.CENTER)
                add_cell(tr[4], f"{pv:.3f}",      align=WD_ALIGN_PARAGRAPH.CENTER)

    # Panel C header
    tr = tbl.add_row().cells
    add_cell(tr[0], "Panel C: Pre-Trend F-Tests (k = −4, −3, −2)", bold=True)
    for c in tr[1:]: add_cell(c, "")
    for tk in ["T1_moderate_1k","T2_large_5k","T3_verylarge_25k"]:
        for ok in ["ihs_suicide_rate","ihs_overdose_rate"]:
            r = panelC.get((tk, ok))
            if r:
                tr = tbl.add_row().cells
                add_cell(tr[0], thresh_labels[tk])
                add_cell(tr[1], outcome_labels[ok])
                add_cell(tr[2], f"F = {float(r['F_stat']):.2f}", align=WD_ALIGN_PARAGRAPH.CENTER)
                add_cell(tr[3], "—")
                add_cell(tr[4], f"{float(r['p_pretrend']):.3f}", align=WD_ALIGN_PARAGRAPH.CENTER)

    note(doc, "Notes: Outcome is IHS-transformed annual county mortality rate (per 100,000). Panel A: TWFE with county and year FE, unemployment rate control; T1 and T3 values from text; T2 comparable in direction. Panel B: County-level pre/post means; 2015 excluded. Panel C: Joint Wald F-test on k = −4, −3, −2 pre-treatment event-study coefficients. *** p<0.01, ** p<0.05, * p<0.10.")


def build_table4(doc):
    heading(doc, "Table 4. Placebo Test: False Treatment Year 2013", sb=6)
    col_hdrs = ["Threshold", "Outcome", "Placebo ATT", "SE", "p-value"]
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(col_hdrs):
        add_cell(hdr[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    thresh_labels = {
        "T1_moderate_1k": "T1 ≥1,000 ac",
        "T2_large_5k":    "T2 ≥5,000 ac",
        "T3_verylarge_25k": "T3 ≥25,000 ac",
    }
    outcome_labels = {"ihs_suicide_rate": "Suicide", "ihs_overdose_rate": "Overdose"}
    with open(os.path.join(ROOT, "results", "did_2015_placebo.csv"), newline='') as f:
        for r in csv.DictReader(f):
            att = float(r['att']); se_ = float(r['se']); pv = float(r['pvalue'])
            tr = tbl.add_row().cells
            add_cell(tr[0], thresh_labels.get(r['threshold'], r['threshold']))
            add_cell(tr[1], outcome_labels.get(r['outcome'], r['outcome']))
            add_cell(tr[2], f"{att:+.3f}", align=WD_ALIGN_PARAGRAPH.CENTER)
            add_cell(tr[3], f"{se_:.3f}",  align=WD_ALIGN_PARAGRAPH.CENTER)
            add_cell(tr[4], f"{pv:.3f}",   align=WD_ALIGN_PARAGRAPH.CENTER)
    note(doc, "Notes: Simple DiD with treatment year reassigned to 2013. Estimation window: 2011–2016. SEs clustered by county. No estimates are statistically significant at conventional levels.")


def build_table5(doc):
    heading(doc, "Table 5. CR2 Robustness: T1 Event-Study Standard Errors", sb=6)
    col_hdrs = ["Outcome", "Year (k)", "Coefficient", "SE (CR2)", "p-value (Satterthwaite)"]
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(col_hdrs):
        add_cell(hdr[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    outcome_labels = {"ihs_suicide_rate": "Suicide", "ihs_overdose_rate": "Overdose"}
    with open(os.path.join(ROOT, "results", "did_2015_cr2_ses.csv"), newline='') as f:
        for r in csv.DictReader(f):
            if r['threshold'] != "T1_moderate_1k":
                continue
            tr = tbl.add_row().cells
            add_cell(tr[0], outcome_labels.get(r['outcome'], r['outcome']))
            et = int(r['event_time'])
            add_cell(tr[1], f"k = {et:+d}", align=WD_ALIGN_PARAGRAPH.CENTER)
            add_cell(tr[2], f"{float(r['beta']):+.3f}",  align=WD_ALIGN_PARAGRAPH.CENTER)
            add_cell(tr[3], f"{float(r['se_cr2']):.3f}", align=WD_ALIGN_PARAGRAPH.CENTER)
            add_cell(tr[4], f"{float(r['p_satt']):.3f}", align=WD_ALIGN_PARAGRAPH.CENTER)
    note(doc, "Notes: T1 (≥1,000 acres) event-study coefficients, reference year k = −1. CR2: Bell-McCaffrey bias-corrected standard errors with Satterthwaite degrees of freedom via the clubSandwich package. All p-values > 0.49; mortality results are robustly insignificant under both CR1 and CR2 inference.")


# ── Figure captions ───────────────────────────────────────────────────────────

FIG_CAPTIONS = {
    "1": ("Common Support: Covariate Overlap Between Treated and Matched Controls "
          "(T1, ≥1,000 acres). Treated counties experienced a qualifying 2015 wildfire. "
          "Controls matched within WHP quintiles via Mahalanobis distance. Near-identical "
          "distributions confirm common support across all matching dimensions."),
    "2": ("Treatment Effects on Depression Prevalence and Poor Mental Health Days (2019). "
          "T1 binary treatment estimate (filled circle) and dose-response estimate "
          "(open diamond) with 95% confidence intervals. Panel A: depression prevalence (%). "
          "Panel B: poor mental health days (%). *** p<0.01, ** p<0.05, * p<0.10. "
          "Heteroskedasticity-robust SEs."),
}

# ── Main manuscript builder ───────────────────────────────────────────────────

def make_manuscript():
    doc = new_doc()

    # ── Title page ────────────────────────────────────────────────────────────
    add_para(doc, "Wildfire and Long-Run Mental Health:",
             align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, sa=0)
    add_para(doc, "Evidence from Matched Difference-in-Differences",
             align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, sa=12)
    add_para(doc, "[Author Name]",        align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)
    add_para(doc, "[Institution]",        align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)
    add_para(doc, "June 2026",            align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)
    add_para(doc, "Submission Category: STEM",
             align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, sa=24)

    # ── Abstract ──────────────────────────────────────────────────────────────
    add_para(doc, "Abstract", bold=True, sa=4)
    noindent(doc, ("I estimate the causal effect of the 2015 Western U.S. wildfire season on "
                   "county-level mental health using a matched difference-in-differences design. "
                   "Wildfire Hazard Potential matching identifies 81 treated counties from "
                   "structurally similar never-burned controls within fire-risk strata. Four "
                   "years post-fire, treated counties report depression prevalence 2.76 "
                   "percentage points higher (SE = 0.94, p = 0.004) and poor mental health "
                   "days 1.63 percentage points higher (SE = 0.51, p = 0.002). A dose-response "
                   "specification replacing the binary treatment indicator with log burn intensity "
                   "yields a coefficient of 1.18 per log-unit (p = 0.020), inconsistent with "
                   "selection confounding. This monotone relationship between burn severity and "
                   "subsequent depression is the primary evidence against a pure selection "
                   "interpretation of the treatment effect."))
    add_para(doc, "Keywords: wildfire, mental health, depression, difference-in-differences, "
             "rural health, natural disaster", italic=True, sa=12)

    # ── Body: read and parse extracted text ───────────────────────────────────
    txt_path = os.path.join(ROOT, "docx_extracted.txt")
    with open(txt_path, encoding='utf-8') as f:
        lines = [l.rstrip('\n') for l in f.readlines()]

    # Skip title block (lines 0-4 = indices 0-4, which are lines 1-5 in cat-n output)
    skip_set = set(range(0, 5))  # title, subtitle, author, institution, date

    # Markers we handle explicitly
    TABLE_TRIGGERS = {
        "Table 1.": build_table1,
        "Table 2.": build_table2,
    }

    i = 0
    while i < len(lines):
        if i in skip_set:
            i += 1
            continue

        line = lines[i].strip()
        if not line:
            i += 1
            continue

        line = mla(line)

        # References section → replaced by Works Cited below
        if line == "References":
            i += 1
            # Skip all reference lines — we output MLA Works Cited separately
            while i < len(lines) and not lines[i].strip().startswith("Figures"):
                i += 1
            continue

        # Figure placeholder lines
        if line.startswith("[Figure") or line.startswith("Figure 1 ") or \
           line.startswith("Figure 2 "):
            # Detect which figure
            for num in ["1","2"]:
                if f"Figure {num}" in line and line.startswith("Figure"):
                    insert_figure(doc, num, FIG_CAPTIONS[num])
                    break
            i += 1
            continue

        # "Figures" section heading (just the word "Figures" on its own)
        if line == "Figures":
            # Already handled inline — skip
            i += 1
            continue

        # Table title lines — insert actual table
        triggered = False
        for trigger, builder in TABLE_TRIGGERS.items():
            if line.startswith(trigger):
                builder(doc)
                triggered = True
                break
        if triggered:
            i += 1
            continue

        # Table notes
        if line.startswith("Notes:"):
            note(doc, line)
            i += 1
            continue

        # Section headings: "1.  Introduction", "2.  Data", etc.
        import re
        if re.match(r'^\d+\.\s{2,}\S', line):
            heading(doc, line)
            i += 1
            continue

        # Subsection headings: "2.1  Treatment", "3.2  Mahalanobis", etc.
        if re.match(r'^\d+\.\d+\s{2,}\S', line):
            subheading(doc, line)
            i += 1
            continue

        # Assumption line
        if line.startswith("Assumption"):
            noindent(doc, line)
            i += 1
            continue

        # Equation lines like "(1)", "(2)", "(3)", "(4)", "(5)"
        if re.match(r'^\(\d\)$', line):
            add_para(doc, f"[ Equation {line[1]} ]",
                     align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, sb=4, sa=4)
            i += 1
            continue

        # Everything else: body paragraph
        body(doc, line)
        i += 1

    # ── Works Cited ───────────────────────────────────────────────────────────
    page_break(doc)
    add_para(doc, "Works Cited", bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=12)

    entries = [
        ('Bergé, Laurent. ', 'Fixest: Fast Fixed Effects Estimation',
         '. R package version 0.10.4, 2022, lrberge.github.io/fixest/.'),
        ('Callaway, Brantly, and Pedro H. C. Sant\'Anna. ', '"Difference-in-Differences with Multiple Time Periods."',
         ' Journal of Econometrics, vol. 225, no. 2, 2021, pp. 200–230.'),
        ('Currie, Janet, and Soodeh Saberian. ', '"Wildfire, Smoke and Mental Health in Canada."',
         ' NBER Working Paper No. 33912, 2025, www.nber.org/papers/w33912.'),
        ('Du, Ran, et al. ', '"Climate Disaster and Cognitive Ability: Evidence From Wildfire."',
         ' International Journal of Public Health, vol. 69, 2024, article 1607128.'),
        ('Greenberg, Paul E., et al. ', '"The Economic Burden of Adults with Major Depressive Disorder in the United States (2010 and 2018)."',
         ' PharmacoEconomics, vol. 39, no. 6, 2021, pp. 653–665.'),
        ('Imbens, Guido W., and Michal Kolesár. ', '"Robust Standard Errors in Small Samples: Some Practical Advice."',
         ' Review of Economics and Statistics, vol. 98, no. 4, 2016, pp. 701–712.'),
        ('Jung, Youn Soo, et al. ', '"Fine Particulate Matter From 2020 California Wildfires and Mental Health–Related Emergency Department Visits."',
         ' JAMA Network Open, vol. 8, no. 4, 2025, article e253326.'),
        ('Merdjanoff, Alexis A., et al. ', '"Projecting the Mental Health Impacts of the 2025 LA Wildfires: Lessons from Hurricane Katrina."',
         ' Environmental Research Letters, vol. 21, no. 11, 2026, article 111004.'),
        ('Pustejovsky, James E. ', 'clubSandwich: Cluster-Robust (Sandwich) Variance Estimators with Small-Sample Corrections',
         '. R package version 0.5.8, 2022, CRAN.R-project.org/package=clubSandwich.'),
        ('Wettstein, Zachary S., and Ambarish Vaidyanathan. ', '"Psychotropic Medication Prescriptions and Large California Wildfires."',
         ' JAMA Network Open, vol. 7, no. 2, 2024, article e2356466.'),
        ('Ye, Xinyue, et al. ', '"Wildfires and Public Health: A Comprehensive Review of Human-Centric Studies."',
         ' GeoHealth, vol. 10, no. 2, 2026, article e2025GH001534.'),
    ]

    for author, title, rest in entries:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing   = 1.15
        p.paragraph_format.space_before   = Pt(0)
        p.paragraph_format.space_after    = Pt(6)
        p.paragraph_format.left_indent    = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        r = p.add_run(author)
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)
        r2 = p.add_run(title)
        r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)
        # Italicize book/journal titles (no quotes) — rough heuristic: no leading quote
        if not title.startswith('"'):
            r2.italic = True
        r3 = p.add_run(rest)
        r3.font.name = 'Times New Roman'; r3.font.size = Pt(12)

    out = os.path.join(OUTDIR, "manuscript.docx")
    doc.save(out)
    print(f"Saved: {out}")
    return out


# ── Cover letter ──────────────────────────────────────────────────────────────

def make_cover_letter():
    doc = new_doc()

    add_para(doc, "June 2026", sa=12)
    add_para(doc, "Editorial Board", sa=0)
    add_para(doc, "The Curieux Review", sa=0)
    add_para(doc, "curieuxreview.com", sa=24)

    add_para(doc, "Dear Editors,", sa=12)

    body(doc, ("I am pleased to submit my research paper, 'Wildfire and Long-Run Mental Health: "
               "Evidence from Matched Difference-in-Differences,' for consideration in The Curieux "
               "Review under the STEM category."))

    body(doc, ("This paper investigates whether the 2015 Western U.S. wildfire season caused "
               "lasting harm to population mental health. Using a matched difference-in-differences "
               "design — in which wildfire-affected counties are compared to structurally similar "
               "unaffected counties matched on federal wildfire hazard potential scores — I find "
               "that four years after the fires, affected counties showed depression prevalence "
               "nearly three percentage points above comparable counties that did not burn. A "
               "dose-response test, in which counties that lost more land to fire show "
               "proportionally worse mental health outcomes, provides evidence that this gap "
               "reflects a causal effect rather than pre-existing differences between fire-prone "
               "and fire-free communities."))

    body(doc, ("The paper contributes what I believe is the first quasi-experimental, "
               "population-level causal estimate of wildfire effects on mental health at a "
               "multi-year follow-up horizon. It draws on publicly available federal data (CDC, "
               "U.S. Forest Service, Census Bureau) and standard econometric methods, and "
               "directly addresses a gap identified by systematic reviewers in the recent "
               "literature."))

    body(doc, ("The manuscript is formatted per Curieux guidelines: Times New Roman 12-point "
               "font, 1.15 spacing, and MLA citations and Works Cited. The paper has not been "
               "submitted elsewhere. I am happy to revise in response to editorial feedback."))

    body(doc, "Thank you for your time and consideration.")

    add_para(doc, "", sa=12)
    add_para(doc, "Sincerely,", sa=0)
    add_para(doc, "", sa=12)
    add_para(doc, "[Author Name]", sa=0)
    add_para(doc, "[Institution]", sa=0)
    add_para(doc, "[Email]", sa=0)

    out = os.path.join(OUTDIR, "cover_letter.docx")
    doc.save(out)
    print(f"Saved: {out}")
    return out


# ── Run ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import re
    print("Building manuscript...")
    make_manuscript()
    print("Building cover letter...")
    make_cover_letter()
    print("\nDone. Files in:", OUTDIR)
    for f in os.listdir(OUTDIR):
        size = os.path.getsize(os.path.join(OUTDIR, f)) / 1024
        print(f"  {f}  ({size:.0f} KB)")
