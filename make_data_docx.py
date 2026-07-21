import sys; sys.path.insert(0, r'C:\Users\chenyon\Research Paper 2026(1)')
from make_equations import *
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(24)
style.paragraph_format.space_after  = Pt(0)
style.paragraph_format.space_before = Pt(0)


def heading(text, level=2, center=False):
    """Section or subsection heading."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    return p


def para(text, indent=True, italic_prefix=None):
    """Body paragraph, optionally with an italic run prefix (for 'Paragraph headings')."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if italic_prefix:
        r0 = p.add_run(italic_prefix + "  ")
        r0.font.name = 'Times New Roman'
        r0.font.size = Pt(12)
        r0.italic = True
        r0.bold   = True
    add_math_runs(p, text)
    return p


def blank():
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# Section 2: Data
# ═══════════════════════════════════════════════════════════════════════════════
heading("2.  Data", center=True)
blank()

# ── 2.1 Treatment ─────────────────────────────────────────────────────────────
heading("2.1  Treatment: Wildfire Occurrence and Severity")
blank()

para(
    "I identify wildfire events using the Monitoring Trends in Burn Severity "
    "(MTBS) fire perimeter database, maintained jointly by the U.S. Geological "
    "Survey and the U.S. Forest Service. MTBS maps the perimeter of every "
    "wildland fire affecting at least 1,000 acres in the Western United States "
    "or 500 acres in the Eastern United States, using Landsat satellite imagery "
    "to delineate burned area boundaries. The database covers all fire seasons "
    "from 1984 through 2022 and reports each event's fire type (wildfire, "
    "prescribed burn, or wildland fire use), ignition year, and total acres "
    "burned. I restrict to events classified as wildfire, excluding prescribed "
    "burns. Prescribed burns are planned events that do not generate the "
    "psychological trauma, unplanned displacement, and acute financial loss "
    "associated with wildfire exposure."
)

para(
    "I assign fires to counties using spatial intersection of MTBS burn "
    "perimeter polygons with 2020 TIGER/Line county shapefiles (U.S. Census "
    "Bureau). County land area in acres is computed from the TIGER/Line ALAND "
    "field (converted from square meters), and treatment intensity is expressed "
    "as the percent of a county's land area burned in 2015:"
)

insert_eq(doc, eq_pctburned(), "1")

para(
    "All regressions "
    "project county shapes onto the U.S. National Atlas Equal Area projection "
    "(EPSG:5070) to ensure area calculations are not distorted by the curvature "
    "of the Earth."
)

# ── 2.2 Mental Health Outcomes ────────────────────────────────────────────────
blank()
heading("2.2  Mental Health Outcomes")
blank()

para(
    "County-year mortality counts are from CDC WONDER, an online query system "
    "that provides restricted access to compressed mortality files derived from "
    "death certificates. I retrieve annual counts of deaths coded to ICD-10 "
    "chapter XX (intentional self-harm, codes X60–X84 and Y87.0) for suicide "
    "and to ICD-10 codes X40–X44 (accidental poisoning by drugs) for drug "
    "overdose mortality. Counts are converted to rates per 100,000 resident "
    "population using intercensal population estimates from the American "
    "Community Survey (ACS).",
    italic_prefix="Suicide and drug overdose mortality (2011–2019)."
)

para(
    "The CDC suppresses WONDER cell counts below ten deaths to protect patient "
    "privacy. In the T1 panel of 93 counties observed over nine years, "
    "approximately 50 percent of county-year cells are suppressed for suicide "
    "and 47 percent for overdose—a direct consequence of the small population "
    "size of many Western counties. Suppressed cells are treated as missing in "
    "the regression, not as zeros; this introduces measurement error and reduces "
    "the effective sample for mortality outcomes. To handle zeros in the "
    "non-suppressed cells, I transform both rates with the inverse hyperbolic "
    "sine (IHS), which approximates ln(2Y) for large values but remains defined "
    "at zero, avoiding the need to drop or impute zero-rate counties:"
)

insert_eq(doc, eq_ihs(), "2")

para(
    "I obtain two survey-based mental health prevalence measures from the CDC "
    "PLACES database (formerly 500 Cities), which reports model-based estimates "
    "of health behavior and outcome prevalence at the county level. "
    "Depression prevalence is the estimated percentage of adults who report ever having "
    "been told by a doctor that they have depressive disorder. "
    "Poor mental health days is the estimated percentage of adults who report 14 or "
    "more poor mental health days in the past 30 days. Both measures are derived "
    "from the Behavioral Risk Factor Surveillance System (BRFSS) using small-area "
    "estimation methods that combine survey responses with local demographic "
    "predictors.",
    italic_prefix="Depression prevalence and poor mental health days (2019)."
)

para(
    "CDC PLACES provides county-level estimates beginning with 2019. Because "
    "2020 is excluded from the analysis to avoid COVID-19 confounding, I use the "
    "2019 wave as a four-year post-treatment cross-section. This cross-sectional "
    "structure is the binding constraint on the depression analysis: I cannot "
    "construct a pre-treatment depression rate from PLACES data and therefore "
    "cannot run a panel DiD for depression. Instead, I estimate cross-sectional "
    "OLS regressions that control for 2014 baseline covariates as proxies for "
    "the pre-treatment mental health environment."
)

# ── 2.3 Covariates ────────────────────────────────────────────────────────────
blank()
heading("2.3  Covariates")
blank()

para(
    "Structural wildfire risk is measured using the Wildfire Risk to Communities "
    "(WRC) county-level dataset published by the U.S. Forest Service. The "
    "primary matching variable, BP_NATIONAL_RANK, is the national percentile "
    "rank of each county's mean annual burn probability as simulated by FSim, a "
    "physics-based large-fire simulation model that uses LANDFIRE vegetation and "
    "fuel data together with historical weather statistics. Because FSim simulates "
    "fire behavior from landscape characteristics rather than historical fire "
    "occurrence, BP_NATIONAL_RANK is plausibly exogenous to realized wildfire "
    "events and serves as a pre-treatment measure of structural fire risk. A "
    "second WRC variable, BUILDINGS_FRACTION_DE, measures the share of county "
    "buildings in the Direct Exposure WHP zone—the highest-risk category—and "
    "enters the matching procedure alongside the rank measure.",
    italic_prefix="Wildfire Hazard Potential (WHP)."
)

para(
    "The 2013 Rural-Urban Continuum Codes (RUCC) published by the USDA Economic "
    "Research Service classify counties on a 1–9 scale from large metropolitan "
    "areas (RUCC 1–3) to completely rural counties (RUCC 7–9). I use the 2013 "
    "vintage as a pre-treatment measure of county rurality and include it as both "
    "a matching variable and a heterogeneity dimension. For the rurality subgroup "
    "analysis in Section 5, I collapse RUCC into a binary indicator distinguishing "
    "metropolitan counties (RUCC 1–3) from non-metropolitan counties (RUCC 4–9).",
    italic_prefix="Rurality."
)

para(
    "Population and median household income are drawn from the ACS 1-year "
    "estimates for 2014, the year immediately before treatment. Annual county "
    "unemployment rates (2011–2019) come from the Bureau of Labor Statistics "
    "Local Area Unemployment Statistics (BLS LAUS) program. The HRSA Health "
    "Professional Shortage Area (HPSA) designation is a binary indicator that "
    "equals one if the county is designated a mental health HPSA in any year "
    "during the study period, capturing structural undersupply of mental health "
    "providers relative to county need.",
    italic_prefix="Socioeconomic covariates."
)

# ── 2.4 Sample Construction and Descriptive Statistics ───────────────────────
blank()
heading("2.4  Sample Construction and Descriptive Statistics")
blank()

para(
    "The analysis covers eleven Western states: Arizona, California, Colorado, "
    "Idaho, Montana, Nevada, New Mexico, Oregon, Utah, Washington, and Wyoming. "
    "I apply a minimum population threshold of 10,000 residents (2014 ACS) to "
    "exclude counties in which virtually all mortality cells would be suppressed "
    "in CDC WONDER."
)

para(
    "A county is treated if it experienced its first MTBS wildfire of qualifying "
    "size in calendar year 2015 and had no such qualifying fire in the prior study "
    "period. Control counties are those with no qualifying wildfire in any year "
    "from 2015 through 2019 (true never-treated counties). Counties that "
    "experienced qualifying fires only in 2016–2019 are excluded from both groups "
    "to keep controls clean. This single-cohort restriction—rather than a "
    "staggered DiD across 2015–2019—is motivated by the density of wildfire "
    "activity in the Western United States: by 2019, nearly all high-WHP counties "
    "had experienced at least one qualifying fire, leaving an insufficient "
    "never-treated control pool for a staggered design.",
    italic_prefix="Treatment assignment."
)

para(
    "Three fire-size thresholds define three separate treated samples. The "
    "primary threshold (T1, ≥1,000 acres) yields 94 treated counties, of which "
    "81 appear in the matched panel after merging with the base outcome data. The "
    "large-fire threshold (T2, ≥5,000 acres) yields 51 treated counties; the "
    "very-large-fire threshold (T3, ≥25,000 acres) yields 19. Across the three "
    "thresholds, mean county land area burned in 2015 ranges from 6.7 percent "
    "(T1) to 24.0 percent (T3).",
    italic_prefix="Treated samples."
)

para(
    "The panel spans 2011–2019, providing four pre-treatment years and four "
    "post-treatment years, with 2015 as the treatment year. The year 2020 is "
    "excluded entirely to avoid COVID-19 confounding; post-2020 data are not used "
    "because the depression cross-section is pinned to 2019.",
    italic_prefix="Study window."
)

para(
    "Table 1 presents summary statistics for the primary T1 matched panel. The "
    "93 counties are observed over nine years, yielding 837 county-year "
    "observations. The sample is predominantly non-metropolitan: mean RUCC is 5.4 "
    "among treated counties and 5.3 among controls. Mean WHP national rank is "
    "0.836 for treated and 0.729 for controls, reflecting within-quintile matching "
    "that equalizes structural fire risk while allowing some residual rank "
    "difference across quintiles. Among non-suppressed observations, mean suicide "
    "rates are 20.2 per 100,000 (SD = 9.7) and mean overdose rates are 22.9 per "
    "100,000 (SD = 14.6)—both elevated relative to the national average, "
    "consistent with the literature on rural Western mortality. Depression "
    "prevalence in 2019 is 21.2 percent among treated counties and 18.0 percent "
    "among controls; the corresponding figures for poor mental health days are "
    "15.1 and 13.2 percent. These raw differences—3.2 and 1.9 percentage "
    "points—are slightly larger than the regression-adjusted estimates in "
    "Section 5, which condition on baseline covariates.",
    italic_prefix="Descriptive statistics."
)

para("[TABLE 1 ABOUT HERE]", indent=False)

out_path = r'C:\Users\chenyon\Research Paper 2026(1)\section_data.docx'
doc.save(out_path)
print(f"Saved: {out_path}")
