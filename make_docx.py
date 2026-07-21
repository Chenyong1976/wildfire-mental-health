from docx import Document
from docx.shared import Pt, Inches

doc = Document()

# Page margins: 1 inch all around
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# Default style: Times New Roman 12pt, double-spaced
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(24)
style.paragraph_format.space_after  = Pt(0)
style.paragraph_format.space_before = Pt(0)


def add_para(text, first_indent=True, bold=False, center=False):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


# ── Title block ───────────────────────────────────────────────────────────────
from docx.enum.text import WD_ALIGN_PARAGRAPH

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.line_spacing = Pt(24)
title_p.paragraph_format.space_after  = Pt(0)
tr = title_p.add_run(
    "Wildfire and Long-Run Mental Health:\n"
    "Evidence from Matched Difference-in-Differences"
)
tr.font.name = 'Times New Roman'
tr.font.size = Pt(12)
tr.bold = True

doc.add_paragraph()  # blank line

author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_p.paragraph_format.line_spacing = Pt(24)
author_p.paragraph_format.space_after  = Pt(0)
ar = author_p.add_run("[Author Name]\n[Institution]\nJune 2026")
ar.font.name = 'Times New Roman'
ar.font.size = Pt(12)

doc.add_paragraph()  # blank line

# ── Section heading ───────────────────────────────────────────────────────────
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
h.paragraph_format.line_spacing = Pt(24)
h.paragraph_format.space_after  = Pt(0)
hr = h.add_run("1.  Introduction")
hr.font.name = 'Times New Roman'
hr.font.size = Pt(12)
hr.bold = True

doc.add_paragraph()  # blank line

# ── Body paragraphs ───────────────────────────────────────────────────────────
paragraphs = [
    (
        "Wildfire is the fastest-growing natural disaster in the United States. "
        "Between 2000 and 2023, the annual area burned in the Western U.S. more than "
        "doubled. In 2015—the most destructive fire season on record at the "
        "time—fires swept through twelve Western states, scorching millions of acres, "
        "forcing mass evacuations, and filling regional air sheds with hazardous smoke "
        "for months. The physical health consequences are well established: wildfire smoke "
        "raises hospitalizations for respiratory and cardiovascular disease, and direct "
        "fire exposure kills. The mental health consequences are far less understood. "
        "Wildfire exposure creates plausible pathways to depression, anxiety, suicidal "
        "ideation, and substance use disorders—through the neuroinflammatory effects "
        "of fine particulate matter, the trauma of forced displacement, the financial "
        "stress of property loss, and the anxiety amplified by media coverage. Despite "
        "these mechanisms, rigorous long-run causal evidence on whether wildfires worsen "
        "population mental health remains absent from the literature."
    ),
    (
        "Counties that experience large wildfires are not randomly selected—they are "
        "systematically more rural, drier, more economically marginal, and farther from "
        "urban healthcare markets than counties that do not burn. If high-fire counties "
        "were already on worsening mental health trajectories before any fire occurred, "
        "a naive comparison overstates the causal effect. The existing literature relies "
        "almost entirely on cross-sectional and interrupted time-series designs. Neither "
        "design can rule out this confound. No published study has applied a "
        "quasi-experimental design to U.S. wildfire mental health outcomes at the "
        "population level."
    ),
    (
        "This paper fills that gap. I estimate the causal effect of the 2015 Western "
        "U.S. wildfire season on county-level depression prevalence, poor mental health "
        "days, suicide rates, and drug overdose mortality using a matched "
        "difference-in-differences design. Treatment counties are Western U.S. counties "
        "that experienced wildfires of at least 1,000 acres in 2015. Control counties "
        "are Western counties with no qualifying wildfire from 2015 through 2019, matched "
        "to treated counties on structural wildfire hazard potential (WHP), rurality, "
        "building exposure density, pre-treatment outcomes, population, and income. "
        "By 2019, wildfire-affected counties report depression prevalence 2.8 percentage "
        "points higher (SE = 0.94, p = 0.004) and poor mental health days 1.6 percentage "
        "points higher (SE = 0.51, p = 0.002) than matched control counties—four "
        "years after the fires, with pre-treatment trends closely parallel between "
        "the two groups."
    ),
    (
        "Both effects are economically meaningful. Among control counties, mean depression "
        "prevalence is 18.0 percent. As such, the 2.8 percentage point "
        "estimate implies a 15 percent increase relative to baseline—comparable to "
        "the depression effect of large-scale unemployment shocks documented in the labor "
        "economics literature. To test for a causal gradient, I replace the binary fire "
        "indicator with the log of the percent of county land area burned. Counties that "
        "burned more land show proportionally worse depression by 2019 "
        "(β = 1.18 per log-unit of land burned, SE = 0.50, p = 0.020). This monotone "
        "relationship is inconsistent with the leading selection confound: a selection "
        "story predicts that fire-prone counties have worse mental health regardless of "
        "fire severity, but predicts no graded response to how much land actually burned. "
        "For suicide and overdose mortality, event-study estimates are consistently "
        "positive but imprecise. These outcomes are measured annually in the panel and "
        "are affected by the small matched control pool of twelve counties. One exception: "
        "among non-metropolitan counties, suicide rates are elevated four years post-fire "
        "(β = 0.241, SE = 0.113, p = 0.038), consistent with the hypothesis that "
        "rural communities face greater and longer-lasting mental health consequences "
        "from wildfire exposure."
    ),
    (
        "The identification strategy rests on the WHP-matched control group. "
        "WHP—a measure derived from the U.S. Forest Service's FSim large-fire "
        "simulation system—captures a county's structural wildfire risk based on "
        "vegetation, topography, and historical weather, independent of whether fires "
        "have occurred. Restricting the control group to counties with similar WHP "
        "converts the identifying assumption from ‘fire-prone and non-fire-prone "
        "counties have parallel mental health trends’—which is difficult to "
        "defend—to ‘among counties with equal structural fire risk, those that "
        "experienced a fire and those that did not would have followed parallel trends "
        "absent the fire.’ Joint pre-trend tests for both mortality outcomes are "
        "clean at the primary threshold (F = 0.53, p = 0.663 for suicide; F = 0.51, "
        "p = 0.679 for overdose). A placebo test assigning a fake treatment year of 2013 "
        "to the same counties produces null results for both suicide "
        "(ATT = 0.160, p = 0.399) and overdose (ATT = 0.140, p = 0.488)."
    ),
    (
        "Three recent papers provide the strongest evidence on wildfire and mental health. "
        "Jung et al. (2025) analyze 86,609 emergency department visits "
        "during the 2020 California wildfire season, isolating wildfire-specific "
        "PM₂.₅ from total ambient PM₂.₅ using satellite-derived smoke "
        "estimates; a 10 μg/m³ increase in wildfire-specific PM₂.₅ "
        "raises mental health ED visits by 8 percent and mood-affective disorder visits "
        "by 29 percent, with effects persisting for up to seven days. Wettstein and "
        "Vaidyanathan (2024) use prescription records from 7.1 million commercially "
        "insured adults across 25 California wildfire events between 2011 and 2018, "
        "documenting 4–6 percent increases in antidepressant, anxiolytic, and mood "
        "stabilizer prescriptions in the six weeks following fire onset. Currie and "
        "Saberian (2025) use Canadian administrative hospitalization data from 2006–2018 "
        "and find that four independent pathways—local PM₂.₅, evacuation "
        "orders, property damage costs, and media-mediated anxiety—each independently "
        "raise mental health hospitalizations. Beyond the mental health literature, Du et al. "
        "(2024) use daily wind direction as an instrument for smoke exposure in China, "
        "finding that a 10-unit increase in upwind wildfires reduces cognitive test "
        "scores by 0.24 standard deviations, providing credible causal evidence for "
        "the neurological mechanism linking wildfire smoke to brain function."
    ),
    (
        "Three limitations leave the causal and long-run picture incomplete. "
        "First, existing studies measure outcomes over short windows—at most six "
        "weeks—leaving effects at multi-year horizons undocumented. If the welfare "
        "costs of wildfire include sustained depression and elevated suicide rates, "
        "short-window studies will understate total harm. Ye et al. (2026), in a "
        "systematic review of 139 wildfire-health studies, identifies the absence of "
        "long-run outcome data as the most consequential gap; Merdjanoff et al. (2026) "
        "draw on Hurricane Katrina research to project the mental health trajectory for "
        "the 2025 LA wildfires, a method made necessary by the same absence of multi-year "
        "wildfire follow-up data. Second, the strongest U.S. study—Jung et al. "
        "(2025)—uses a 2020 study period during which COVID-19 independently drove "
        "increases in mental health ED visits, disrupted healthcare utilization, and "
        "caused economic shocks. These effects are observationally indistinguishable from "
        "wildfire-driven effects in the study design, so effect sizes cannot be cleanly "
        "attributed to wildfire exposure. Third, no design controls for the selection of "
        "vulnerable populations into fire-prone areas—existing comparison groups are "
        "chosen for geographic proximity, not structural wildfire risk equivalence."
    ),
    (
        "This paper addresses each limitation. It provides the first quasi-experimental "
        "causal estimates of wildfire effects on mental health using a matched DiD that "
        "equates treated and control counties on structural fire risk, the first outcome "
        "measurements at a four-year follow-up horizon, and the first dose-response "
        "analysis linking fire severity to subsequent depression prevalence. By excluding "
        "the pandemic year 2020 and drawing depression outcomes from 2019, the design "
        "avoids the COVID-19 confound that complicates interpretation of the most-cited "
        "recent estimates. The rurality heterogeneity result adds a direct test of the "
        "mechanism hypothesis that has organized prior work: if mental health effects are "
        "larger in non-metropolitan counties with fewer providers, mental health care "
        "access—rather than intrinsic rural vulnerability—may be the operative "
        "channel, with direct implications for how post-disaster resources are allocated."
    ),
    (
        "The paper proceeds as follows. Section 2 describes the data sources and "
        "treatment construction. Section 3 presents the WHP-based matching procedure "
        "and documents pre-treatment balance. Section 4 lays out the event-study and "
        "cross-sectional specifications. Section 5 presents the main results. "
        "Section 6 reports robustness checks, including alternative fire-size "
        "thresholds, collapsed 2×2 DiD estimates, and CR2 bias-corrected standard "
        "errors. Section 7 concludes with implications for disaster mental health "
        "policy."
    ),
]

for text in paragraphs:
    add_para(text)

out_path = r'C:\Users\chenyon\Research Paper 2026(1)\introduction.docx'
doc.save(out_path)
print(f"Saved: {out_path}")
