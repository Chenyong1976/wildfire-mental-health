import sys; sys.path.insert(0, r'C:\Users\chenyon\Research Paper 2026(1)')
from make_equations import *
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(24)
style.paragraph_format.space_after  = Pt(0)
style.paragraph_format.space_before = Pt(0)


def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True
    return p


def para(text, indent=True, italic_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if italic_prefix:
        r0 = p.add_run(italic_prefix + "  ")
        r0.font.name = 'Times New Roman'
        r0.font.size = Pt(12)
        r0.italic = True
        r0.bold   = True
    add_math_runs(p, text)
    return p



def blank():
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_after  = Pt(0)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3
# ═══════════════════════════════════════════════════════════════════════════════
h1("3.  Matching Procedure")
blank()

para(
    "Standard two-way fixed effects estimation requires that treated and control "
    "counties would have followed parallel mental health trends absent the 2015 "
    "fires. Counties that experience large wildfires are systematically different "
    "from those that do not: they are more rural, more economically marginal, and "
    "more exposed to other environmental stressors. If these structural differences "
    "translate into diverging mental health trajectories—declining provider access "
    "in rural areas, worsening economic conditions in resource-dependent "
    "counties—then a naive comparison conflates causal wildfire effects with "
    "pre-existing divergence. The matching procedure restricts the control group "
    "to structurally similar counties, converting the identifying assumption to a "
    "comparison within fire-risk strata rather than across them."
)

blank()
h2("3.1  WHP Quintile Stratification")
blank()

para(
    "I first assign each Western county to one of five quintiles of "
    "BP_NATIONAL_RANK, the FSim-derived structural burn probability measure "
    "described in Section 2. Control counties can only be matched to treated "
    "counties within the same WHP quintile, ensuring that every treated unit is "
    "compared to counties facing equal structural wildfire risk. Within a WHP "
    "quintile, counties share broadly similar vegetation type, climate zone, "
    "topography, and rural economic profile—the characteristics most likely to "
    "generate correlated mental health trends independent of wildfire occurrence. "
    "Because the WHP raster was last updated with 2014 LANDFIRE data before the "
    "2015 fire season, quintile assignments are pre-treatment and not contaminated "
    "by post-fire vegetation changes."
)

blank()
h2("3.2  Mahalanobis Nearest-Neighbor Matching")
blank()

para(
    "Within each quintile, I match each treated county to up to two control "
    "counties using Mahalanobis distance computed over eight pre-treatment "
    "covariates. The distance between treated county t and candidate control c is:"
)

insert_eq(doc, eq_mahalanobis(), "1")

para(
    "where x is the vector of matching covariates, Σ̂ is the sample "
    "covariance matrix estimated from pooled treated and control observations "
    "within the quintile, and λ = 0.01 × tr(Σ̂) / p is a Tikhonov "
    "regularization term (p = 8 covariates) that prevents near-singular "
    "inversions in quintiles with few observations. The eight matching covariates "
    "are: WHP national rank, rural-urban continuum code (RUCC2013), share of "
    "buildings in the direct exposure zone (BUILDINGS_FRACTION_DE), a binary "
    "indicator for any qualifying wildfire before 2015, 2014 county population, "
    "2014 median household income, and the county-level means of IHS-transformed "
    "suicide and overdose rates over 2011–2014. Including the pre-treatment outcome "
    "means directly in the distance metric operationalizes the parallel-trends "
    "assumption: matched pairs are close not only on structural characteristics "
    "but on the level and trend of the mental health outcomes themselves."
)

blank()
h2("3.3  Geographic Exclusion Constraint")
blank()

para(
    "I impose a 50-kilometer minimum separation between each treated county "
    "centroid and its matched controls. This prevents matching a treated county "
    "to an immediately adjacent county that may share the same fire plume, "
    "evacuation area, or media market—all of which could contaminate the control "
    "group's outcomes. If no control county within the same WHP quintile satisfies "
    "the distance constraint, the constraint is relaxed and the nearest-quintile "
    "control is used as a fallback, which occurs in fewer than 5 percent of "
    "treated counties. Geographic coordinates are computed as county centroids "
    "in the U.S. National Atlas Equal Area projection (EPSG:5070)."
)

blank()
h2("3.4  Balance")
blank()

para(
    "Table 1, Panel A shows that the matching achieves close balance on all "
    "pre-treatment covariates. WHP national rank differs by 0.013 between treated "
    "and control counties (p = 0.70), rural-urban continuum codes differ by 0.21 "
    "(p = 0.81), and income differs by $3,600 (p = 0.34)—none statistically "
    "distinguishable from zero. The pre-treatment mortality rates in Panel B show "
    "larger raw differences (4.4 per 100,000 for suicide, 6.4 for overdose), but "
    "both are statistically insignificant (p > 0.19), reflecting CDC WONDER "
    "suppression-induced selection rather than actual pre-trend divergence."
)

blank()
h2("3.5  Pre-Treatment Parallel Trends")
blank()

para(
    "The parallel trends assumption requires that, absent the 2015 wildfires, "
    "treated and matched control counties would have followed the same trajectory "
    "in mental health outcomes. Formally, for t < 2015:"
)

para(
    "E[Y_{ct}(0) − Y_{c,t−1}(0) | Treated_{c} = 1]  =  "
    "E[Y_{ct}(0) − Y_{c,t−1}(0) | Treated_{c} = 0]",
    indent=False,
    italic_prefix="Assumption (PT)."
)

para(
    "where Y_{ct}(0) is the potential outcome county c would have experienced without "
    "wildfire exposure in year t. The matching procedure is designed to make this "
    "assumption plausible by restricting the control group to counties with equal "
    "structural fire risk, similar rural economic profiles, and close pre-treatment "
    "mental health levels. Because neither assumption can be verified directly, I "
    "test it empirically using the pre-treatment event-study coefficients."
)

para(
    "I estimate the three pre-treatment event-study coefficients (k = −4, −3, −2) "
    "from equation (2) — corresponding to 2011, 2012, and 2013, before any "
    "qualifying wildfire occurred — and test whether they are jointly zero via a "
    "Wald F-test. At the T1 threshold, this joint test yields F = 0.53 (p = 0.663) "
    "for suicide and F = 0.51 (p = 0.679) for overdose, well within conventional "
    "significance thresholds. Individual pre-treatment coefficients are small and "
    "statistically indistinguishable from zero: the largest, at k = −2 for suicide, "
    "is β̂_{k=−2} = 0.487 (SE = 0.431, p = 0.262). The T3 threshold similarly "
    "passes the pre-trend test (F = 0.85, p = 0.478 for suicide). The T2 pre-trend "
    "test is borderline (F = 2.09, p = 0.110 for suicide; F = 2.33, p = 0.083 for "
    "overdose); T2 estimates are therefore interpreted cautiously throughout. "
    "Figures 3 and 4, presented at the end of the paper, plot all pre- and post-treatment "
    "event-study coefficients with 95% confidence intervals. Figure 3 overlays all three "
    "thresholds for each outcome in a single panel; Figure 4 shows the same estimates "
    "separately by fire-size threshold. The flat pre-treatment profile at T1 and T3 provides visual "
    "confirmation of the parallel trends assumption."
)

blank()
h2("3.6  Common Support")
blank()

para(
    "The matching estimator is valid only over the region of common support — "
    "the set of covariate values for which both treated and control counties "
    "exist. Outside that region, counterfactual outcomes for treated counties "
    "must be extrapolated rather than interpolated, which is not guaranteed to "
    "be reliable. Two design features enforce common support in this application."
)

para(
    "First, WHP quintile stratification ensures that every treated county is "
    "compared only to control counties with similar structural fire risk. Because "
    "control counties are drawn from the same quintile as the treated county, "
    "the matching never extrapolates outside the fire-risk distribution of the "
    "control group. Among the 79 T1 treated counties in the matched panel, fewer "
    "than 5 percent required the geographic-constraint fallback (nearest-quintile "
    "rather than same-quintile match), and in all such cases the matched control "
    "lies in an adjacent quintile with a WHP rank difference of less than 0.08."
)

para(
    "Second, I verify covariate overlap directly. The WHP national rank ranges "
    "from 0.208 to 1.000 for treated counties and 0.279 to 0.992 for matched "
    "controls; the ranges are nearly identical. "
    "RUCC scores and pre-treatment mortality rates show similar overlap. "
    "Figure 1 plots the full covariate distributions side by side for treated "
    "and matched control counties. The distributions are nearly indistinguishable "
    "in location and spread, confirming that the matched control group provides "
    "interpolative—not extrapolative—counterfactuals for every treated county."
)

para(
    "Any treated county for which no control county exists in the same WHP "
    "quintile within the 50-kilometer exclusion radius is either assigned a "
    "fallback match (nearest-quintile) or dropped from the analysis if no "
    "control county is within two adjacent quintiles. In practice, no counties "
    "are dropped on this criterion at the T1 threshold. The implication is that "
    "the matched sample covers the full support of the 2015 wildfire-affected "
    "county distribution, and the ATT estimate in Section 5 applies to that "
    "full population of affected counties rather than a trimmed subset."
)

p_cs = doc.add_paragraph()
p_cs.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cs.paragraph_format.line_spacing = Pt(24)
p_cs.paragraph_format.space_after  = Pt(0)
r_cs = p_cs.add_run('[Figure 1 about here]')
r_cs.italic = True
r_cs.font.name = 'Times New Roman'
r_cs.font.size = Pt(12)

blank()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4
# ═══════════════════════════════════════════════════════════════════════════════
h1("4.  Empirical Specifications")
blank()

blank()
h2("4.1  TWFE Event Study (Primary Specification)")
blank()

para(
    "Because all treated counties share a single treatment year (2015), the "
    "identifying variation is the same as in a canonical two-period "
    "difference-in-differences — heterogeneity-robust estimators such as Callaway "
    "and Sant'Anna (2021) collapse to standard TWFE in this single-cohort setting, "
    "so I use TWFE directly. The primary specification is an event-study "
    "regression:"
)

insert_eq(doc, eq_twfe(), "2")

para(
    "where Ỹ_{ct} is the IHS-transformed suicide or overdose mortality rate in "
    "county c in year t; α_{c} and α_{t} are county and year fixed effects; "
    "Treated_{c} equals one for counties that experienced a qualifying 2015 wildfire "
    "and zero for matched controls; and u_{ct} is the county unemployment rate, the "
    "only time-varying covariate included. The coefficients β_{k} trace the "
    "effect of the 2015 wildfire at each relative year k, normalized to zero at "
    "the reference year k = −1 (calendar year 2014). The pre-treatment "
    "coefficients (k = −4, −3, −2) test the parallel-trends assumption via a "
    "joint Wald F-test; the post-treatment coefficients (k = 0, +1, +2, +3, +4) "
    "trace the dynamic effect through four years post-fire. Standard errors are "
    "clustered at the county level. All regressions are estimated in R using the "
    "fixest package (Berge, 2022)."
)

blank()
h2("4.2  Simple DiD (Pooled ATT)")
blank()

para(
    "As a scalar summary of the average post-treatment effect, I estimate:"
)

insert_eq(doc, eq_simple_did(), "3")

para(
    "where Post_{t} = 𝟙[t ≥ 2015]. The coefficient δ_{ATT} is the average "
    "treatment effect on the treated (ATT), pooling across all post-treatment "
    "years. This specification sacrifices the dynamic profile of the event study "
    "in exchange for a single, more precisely estimable summary—useful when "
    "event-study coefficients are individually underpowered, as they are for "
    "suicide and overdose given the large share of suppressed CDC WONDER cells."
)

blank()
h2("4.3  Collapsed 2×2 DiD")
blank()

para(
    "A further aggregation collapses each county to two observations: a pre-period "
    "mean (2011–2014) and a post-period mean (2016–2019), excluding the treatment "
    "year 2015. I then estimate equation (3) on the collapsed two-period panel. "
    "This reduces the influence of annual measurement noise from suppressed "
    "mortality cells and is standard practice for mortality outcomes in small-N "
    "panels where a significant share of cells are censored."
)

blank()
h2("4.4  Cross-Sectional Depression Regressions")
blank()

para(
    "Because CDC PLACES provides depression data for 2019 only, I estimate a "
    "cross-sectional OLS regression:"
)

insert_eq(doc, eq_cs_ols(), "4")

para(
    "where D_{c}^{2019} is the 2019 depression prevalence in county c and x_{c} is a "
    "vector of baseline controls: the 2011–2014 means of IHS suicide and overdose "
    "rates, the 2011–2014 mean unemployment rate, and the 2013 RUCC score. These "
    "pre-period controls proxy for the pre-treatment mental health environment, "
    "approximating the within-DiD comparison that a longitudinal depression measure "
    "would enable directly. I estimate the same regression replacing depression "
    "prevalence with poor mental health days (≥14 days per month) as a second "
    "outcome. The identifying assumption is that, conditional on the 2011–2014 "
    "baseline controls, treated and control counties would have had equal 2019 "
    "depression prevalence absent the fires. Standard errors are heteroskedasticity-robust."
)

blank()
h2("4.5  Dose-Response Specification")
blank()

para(
    "To evaluate whether the depression finding reflects a causal wildfire effect "
    "rather than selection of structurally disadvantaged counties into treatment, I "
    "replace the binary Treated_c indicator with the log of percent county land "
    "area burned in 2015:"
)

insert_eq(doc, eq_dose(), "5")

para(
    "estimated among treated counties only. A selection story predicts a positive "
    "β in equation (4)—fire-prone counties have worse mental health regardless "
    "of fire severity—but predicts no relationship between how much a county "
    "burned and its subsequent depression: β_{D} = 0 in equation (5). A "
    "positive, statistically significant β_{D} is therefore harder to reconcile "
    "with selection and more consistent with a causal effect of fire exposure. "
    "The log(1 + ·) transformation handles the skewed distribution of county-level "
    "burn intensity and is defined for counties with PctBurned_{c} = 0."
)

blank()
h2("4.6  Inference with Few Clusters")
blank()

para(
    "The T1 panel contains 93 county clusters; the T3 panel contains 31. "
    "Cluster-robust standard errors (CR1) are consistent as the number of clusters "
    "grows but may be anti-conservative with fewer than approximately 40 balanced "
    "clusters—a concern especially for T3 and for the rurality subgroup analysis "
    "where treated and control cluster counts are unequal. As a robustness check, "
    "I re-estimate the T1 and T3 TWFE event studies using CR2 (Bell-McCaffrey) "
    "bias-corrected standard errors with Satterthwaite degrees of freedom, "
    "implemented via the clubSandwich package. CR2 applies a finite-sample "
    "correction that accounts for cluster leverage and is the recommended "
    "alternative when the wild cluster bootstrap is unavailable (Imbens and "
    "Kolesar, 2016). Results with CR2 standard errors are reported alongside the "
    "CR1 baseline in Section 6."
)

out = r"C:\Users\chenyon\Research Paper 2026(1)\section_empirical.docx"
doc.save(out)
print(f"Saved: {out}")
