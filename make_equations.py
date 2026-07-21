"""
make_equations.py — OMML math equation helper for python-docx.

Provides reusable builders for native Word math equations via XML injection.
Import with:  from make_equations import *
"""
import re
from lxml import etree
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ── Namespace constants ────────────────────────────────────────────────────────
M  = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def _m(tag):
    return "{%s}%s" % (M, tag)

def _w(tag):
    return "{%s}%s" % (W, tag)


# ── Low-level element builders ─────────────────────────────────────────────────

def r(text, italic=True):
    """Math run — italic by default (variables), non-italic for operators."""
    rPr = etree.Element(_m("rPr"))
    sty = etree.SubElement(rPr, _m("sty"))
    sty.set(_m("val"), "i" if italic else "p")

    run = etree.Element(_m("r"))
    run.append(rPr)
    t = etree.SubElement(run, _m("t"))
    t.text = text
    return run


def op(text):
    """Upright (roman) operator run — for +, −, =, ·, etc."""
    return r(text, italic=False)


def sub(base, *sub_elements):
    """Subscript: base_{sub_elements}  →  <m:sSub>"""
    e = etree.Element(_m("sSub"))
    e_base = etree.SubElement(e, _m("e"))
    e_base.append(base)
    sub_el = etree.SubElement(e, _m("sub"))
    for s in sub_elements:
        sub_el.append(s)
    return e


def sup(base, *sup_elements):
    """Superscript: base^{sup_elements}  →  <m:sSup>"""
    e = etree.Element(_m("sSup"))
    e_base = etree.SubElement(e, _m("e"))
    e_base.append(base)
    sup_el = etree.SubElement(e, _m("sup"))
    for s in sup_elements:
        sup_el.append(s)
    return e


def subsup(base, sub_els, sup_els):
    """Sub+superscript: base_{sub}^{sup}  →  <m:sSubSup>"""
    e = etree.Element(_m("sSubSup"))
    e_base = etree.SubElement(e, _m("e"))
    e_base.append(base)
    sub_el = etree.SubElement(e, _m("sub"))
    if isinstance(sub_els, list):
        for s in sub_els:
            sub_el.append(s)
    else:
        sub_el.append(sub_els)
    sup_el = etree.SubElement(e, _m("sup"))
    if isinstance(sup_els, list):
        for s in sup_els:
            sup_el.append(s)
    else:
        sup_el.append(sup_els)
    return e


def acc(base, char):
    """Accent over base (tilde ~, hat ^)  →  <m:acc>"""
    # Map shorthand chars to Unicode combining characters
    char_map = {
        '~': '̃',   # combining tilde
        '^': '̂',   # combining circumflex
        '̂': '̂',
        '̃': '̃',
    }
    accent_char = char_map.get(char, char)

    accPr = etree.Element(_m("accPr"))
    chr_el = etree.SubElement(accPr, _m("chr"))
    chr_el.set(_m("val"), accent_char)

    e = etree.Element(_m("acc"))
    e.append(accPr)
    e_base = etree.SubElement(e, _m("e"))
    e_base.append(base)
    return e


def dgroup(*body, beg='(', end=')'):
    """Delimited group: beg ... end  →  <m:d>"""
    dPr = etree.Element(_m("dPr"))
    begChr = etree.SubElement(dPr, _m("begChr"))
    begChr.set(_m("val"), beg)
    endChr = etree.SubElement(dPr, _m("endChr"))
    endChr.set(_m("val"), end)

    e = etree.Element(_m("d"))
    e.append(dPr)
    content = etree.SubElement(e, _m("e"))
    for child in body:
        content.append(child)
    return e


def nary_sum(sub_el, sup_el, *body):
    """Summation  →  <m:nary> with chr=∑"""
    naryPr = etree.Element(_m("naryPr"))
    chr_el = etree.SubElement(naryPr, _m("chr"))
    chr_el.set(_m("val"), '∑')   # ∑
    limLoc = etree.SubElement(naryPr, _m("limLoc"))
    limLoc.set(_m("val"), "subSup")

    e = etree.Element(_m("nary"))
    e.append(naryPr)

    sub_wrap = etree.SubElement(e, _m("sub"))
    sub_wrap.append(sub_el)

    sup_wrap = etree.SubElement(e, _m("sup"))
    sup_wrap.append(sup_el)

    body_wrap = etree.SubElement(e, _m("e"))
    for child in body:
        body_wrap.append(child)
    return e


def frac(num, den):
    """Fraction num/den  →  <m:f>"""
    e = etree.Element(_m("f"))
    num_wrap = etree.SubElement(e, _m("num"))
    num_wrap.append(num)
    den_wrap = etree.SubElement(e, _m("den"))
    den_wrap.append(den)
    return e


def rad(*radicand):
    """Square root (degree hidden)  →  <m:rad>"""
    radPr = etree.Element(_m("radPr"))
    degHide = etree.SubElement(radPr, _m("degHide"))
    degHide.set(_m("val"), "1")

    e = etree.Element(_m("rad"))
    e.append(radPr)
    deg_wrap = etree.SubElement(e, _m("deg"))   # empty degree
    body_wrap = etree.SubElement(e, _m("e"))
    for child in radicand:
        body_wrap.append(child)
    return e


def omath(*children):
    """Top-level math container  →  <m:oMath>"""
    e = etree.Element(_m("oMath"))
    for child in children:
        e.append(child)
    return e


# ── In-text math helper ───────────────────────────────────────────────────────

def add_math_runs(p, text, font_size=12, font_name='Times New Roman'):
    """
    Add runs to paragraph p with proper subscript/superscript formatting.
    Use _{...} for subscript and ^{...} for superscript in text.
    All other text is output as regular runs.
    Subscript/superscript runs are set to 8pt (approx 67% of 12pt base).
    """
    SUB_SIZE  = round(font_size * 0.67)   # ≈ 8pt for 12pt text
    parts = re.split(r'(_{[^}]*}|\^{[^}]*})', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('_{') and part.endswith('}'):
            run = p.add_run(part[2:-1])
            run.font.subscript = True
            run.font.size = Pt(SUB_SIZE)
        elif part.startswith('^{') and part.endswith('}'):
            run = p.add_run(part[2:-1])
            run.font.superscript = True
            run.font.size = Pt(SUB_SIZE)
        else:
            run = p.add_run(part)
            run.font.size = Pt(font_size)
        run.font.name = font_name


# ── Paragraph insertion ────────────────────────────────────────────────────────

def insert_eq(doc, omath_el, label):
    """
    Insert a display equation paragraph into doc.

    Layout: centered <m:oMathPara> on one paragraph,
    then a right-aligned label paragraph '(label)'.
    """
    # ── Equation paragraph ────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    # Remove first-line indent
    p.paragraph_format.first_line_indent = Inches(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Build <m:oMathPara>
    oMathPara = etree.Element(_m("oMathPara"))
    oMathParaPr = etree.SubElement(oMathPara, _m("oMathParaPr"))
    jc = etree.SubElement(oMathParaPr, _m("jc"))
    jc.set(_m("val"), "center")
    oMathPara.append(omath_el)

    p._p.append(oMathPara)

    # ── Label paragraph ───────────────────────────────────────────────────────
    lp = doc.add_paragraph()
    lp.paragraph_format.line_spacing = Pt(24)
    lp.paragraph_format.space_after  = Pt(0)
    lp.paragraph_format.space_before = Pt(0)
    lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lr = lp.add_run(f"({label})")
    lr.font.name = 'Times New Roman'
    lr.font.size = Pt(12)

    return p


# ══════════════════════════════════════════════════════════════════════════════
# Named equation builders
# ══════════════════════════════════════════════════════════════════════════════

def eq_pctburned():
    """
    PctBurned_c = AcresBurned2015_c / LandArea_c × 100
    """
    lhs = sub(r('PctBurned'), r('c'))
    rhs_num = sub(r('AcresBurned 2015'), r('c'))
    rhs_den = sub(r('LandArea'), r('c'))
    eq_frac = frac(rhs_num, rhs_den)
    return omath(
        lhs,
        op(' = '),
        eq_frac,
        op(' × 100'),   # × 100
    )


def eq_ihs():
    """
    Ỹ_ct = ln(Y_ct + √(Y²_ct + 1))
    """
    Y_tilde = acc(r('Y'), '~')
    lhs = sub(Y_tilde, r('ct'))

    # √(Y²_ct + 1)
    Y_sq_ct = sub(sup(r('Y'), r('2')), r('ct'))
    sqrt_arg = rad(Y_sq_ct, op(' + 1'))

    # Y_ct + √(...)
    Y_ct = sub(r('Y'), r('ct'))
    inner = dgroup(Y_ct, op(' + '), sqrt_arg)

    return omath(
        lhs,
        op(' = ln'),
        inner,
    )


def eq_mahalanobis():
    """
    d(t, c) = (x_t − x_c)^T (Σ̂ + λI)^{−1} (x_t − x_c)
    """
    x_t = sub(r('x'), r('t'))
    x_c = sub(r('x'), r('c'))

    # (x_t − x_c)
    diff_grp = dgroup(x_t, op(' − '), x_c)

    # (x_t − x_c)^T
    diff_T = sup(diff_grp, r('T'))

    # Σ̂ — use Unicode Sigma with combining hat
    sigma_hat = r('Σ̂')   # Σ + combining circumflex

    # (Σ̂ + λI)
    sig_grp = dgroup(sigma_hat, op(' + λI'))

    # (Σ̂ + λI)^{−1}
    sig_inv = sup(sig_grp, op('−1'))   # −1

    # d(t, c)
    d_tc = r('d')
    args_grp = dgroup(r('t'), op(', '), r('c'))

    return omath(
        d_tc,
        args_grp,
        op(' = '),
        diff_T,
        sig_inv,
        dgroup(x_t, op(' − '), x_c),
    )


def eq_twfe():
    """
    Ỹ_ct = α_c + α_t + Σ_{k=−4, k≠−1}^{+4} β_k · 1[t=2015+k] · Treated_c
           + δ · u_ct + ε_ct
    """
    Y_tilde_ct = sub(acc(r('Y'), '~'), r('ct'))

    alpha_c = sub(r('α'), r('c'))   # α_c
    alpha_t = sub(r('α'), r('t'))   # α_t

    # nary sum limits
    sum_sub = r('k=−4, k≠1')   # k=−4, k≠−1   (U+2212 minus, U+2260 neq)
    sum_sup = r('+4')

    beta_k   = sub(r('β'), r('k'))
    treated_c = sub(r(' Treated'), r('c'))
    indicator = op(' · 1[t=2015+k] · ')   # · 1[...] ·

    sum_body = [beta_k, indicator, treated_c]

    sigma = nary_sum(sum_sub, sum_sup, *sum_body)

    delta    = r('δ')          # δ
    u_ct     = sub(r('u'), r('ct'))
    eps_ct   = sub(r('ε'), r('ct'))   # ε_ct

    return omath(
        Y_tilde_ct,
        op(' = '),
        alpha_c,
        op(' + '),
        alpha_t,
        op(' + '),
        sigma,
        op(' + δ · '),
        u_ct,
        op(' + '),
        eps_ct,
    )


def eq_simple_did():
    """
    Ỹ_ct = α_c + α_t + δ_ATT · Treated_c × Post_t + δ · u_ct + ε_ct
    """
    Y_tilde_ct = sub(acc(r('Y'), '~'), r('ct'))
    alpha_c    = sub(r('α'), r('c'))
    alpha_t    = sub(r('α'), r('t'))
    delta_ATT  = sub(r('δ'), r('ATT'))
    treated_c  = sub(r('Treated'), r('c'))
    post_t     = sub(r('Post'), r('t'))
    u_ct       = sub(r('u'), r('ct'))
    eps_ct     = sub(r('ε'), r('ct'))

    return omath(
        Y_tilde_ct,
        op(' = '),
        alpha_c,
        op(' + '),
        alpha_t,
        op(' + '),
        delta_ATT,
        op(' · '),
        treated_c,
        op(' × '),
        post_t,
        op(' + δ · '),
        u_ct,
        op(' + '),
        eps_ct,
    )


def eq_cs_ols():
    """
    D^{2019}_c = α + β · Treated_c + x′_c γ + ε_c
    """
    D_c_2019  = subsup(r('D'), r('c'), r('2019'))
    beta      = r('β')
    treated_c = sub(r('Treated'), r('c'))
    # x'_c  — prime on x, subscript c
    xprime_c  = sub(sup(r('x'), op('′')), r('c'))   # x′_c  U+2032 prime
    gamma     = r('γ')
    eps_c     = sub(r('ε'), r('c'))

    return omath(
        D_c_2019,
        op(' = α + '),
        beta,
        op(' · '),
        treated_c,
        op(' + '),
        xprime_c,
        gamma,
        op(' + '),
        eps_c,
    )


def eq_dose():
    """
    D^{2019}_c = α + β_D · log(1 + PctBurned_c) + x′_c γ + ε_c
    """
    D_c_2019  = subsup(r('D'), r('c'), r('2019'))
    beta_D    = sub(r('β'), r('D'))
    pct_c     = sub(r('PctBurned'), r('c'))
    log_arg   = dgroup(op('1 + '), pct_c)
    xprime_c  = sub(sup(r('x'), op('′')), r('c'))
    gamma     = r('γ')
    eps_c     = sub(r('ε'), r('c'))

    return omath(
        D_c_2019,
        op(' = α + '),
        beta_D,
        op(' · log'),
        log_arg,
        op(' + '),
        xprime_c,
        gamma,
        op(' + '),
        eps_c,
    )
