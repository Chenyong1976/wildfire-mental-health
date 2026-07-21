"""Generate figures.docx — end-of-paper figures section."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
# Standard page setup (1-inch margins)
section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
for attr in ('left', 'right', 'top', 'bottom'):
    setattr(section, f'{attr}_margin', Inches(1))

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(24)
style.paragraph_format.space_after  = Pt(0)

FIG1 = r"C:\Users\chenyon\Research Paper 2026(1)\fig_eventstudy.png"
FIG2 = r"C:\Users\chenyon\Research Paper 2026(1)\fig_rucc.png"
FIG3 = r"C:\Users\chenyon\Research Paper 2026(1)\fig_common_support.png"


def heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(24)
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)


def caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)


def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.space_after  = Pt(12)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)


def blank():
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_after  = Pt(0)


# Section heading
heading("Figures")
blank()

# Figure 1
p1 = doc.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.space_before = Pt(6)
p1.paragraph_format.space_after  = Pt(0)
p1.add_run().add_picture(FIG1, width=Inches(6.2))

caption("Figure 1. Event-Study Coefficients: Wildfire Effects on Suicide and Overdose Mortality, by Fire-Size Threshold")
note("Notes: Each panel plots TWFE event-study coefficients (filled circles) and 95% confidence intervals from equation (2), with reference year k = −1 (2014) normalized to zero (hollow circle). Outcome is IHS-transformed mortality rate per 100,000. Shaded region = pre-treatment window. Dashed vertical line separates pre- from post-treatment. SEs clustered by county.")
blank()

# Figure 2 on new page
doc.add_page_break()

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(6)
p2.paragraph_format.space_after  = Pt(0)
p2.add_run().add_picture(FIG2, width=Inches(5.8))

caption("Figure 2. Rurality Heterogeneity: Event-Study Coefficients by Metro / Non-Metro Classification (T1, ≥1,000 acres)")
note("Notes: T1 threshold (≥ 1,000 acres). Each panel plots TWFE event-study coefficients from equation (2) estimated separately within metropolitan (RUCC 1–3) and non-metropolitan (RUCC 4–9) county subgroups. Non-metropolitan CI bars clipped at ±2.5 IHS units for readability; true CIs are wider for k = −2 to k = +3. * p < 0.05 at k = +4 (non-metropolitan suicide). SEs clustered by county.")

# Figure 3 on new page
doc.add_page_break()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(6)
p3.paragraph_format.space_after  = Pt(0)
p3.add_run().add_picture(FIG3, width=Inches(6.0))

caption("Figure 3. Common Support: Covariate Overlap Between Treated and Matched Controls (T1, ≥1,000 acres)")
note("Notes: Treated counties are those experiencing a qualifying 2015 wildfire (T1 threshold, ≥ 1,000 acres). Controls are matched within WHP quintiles via Mahalanobis distance. Top-left: WHP national rank distributions; vertical dashed lines mark quintile boundaries (0.2, 0.4, 0.6, 0.8). Top-right: paired WHP ranks, colored by quintile; 45-degree line = perfect match. Bottom panels: RUCC and pre-treatment suicide rate distributions. Near-identical distributions confirm common support across all matching dimensions.")

out = r"C:\Users\chenyon\Research Paper 2026(1)\figures.docx"
doc.save(out)
print(f"Saved: {out}")
