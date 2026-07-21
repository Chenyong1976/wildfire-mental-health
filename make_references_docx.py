"""
Generate references.docx — formatted bibliography to append to combined paper.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(24)
style.paragraph_format.space_after  = Pt(0)
style.paragraph_format.space_before = Pt(0)


def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True


def ref(text):
    """Hanging-indent reference entry."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing    = Pt(24)
    p.paragraph_format.space_after     = Pt(0)
    p.paragraph_format.space_before    = Pt(0)
    p.paragraph_format.left_indent     = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)


def blank():
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_after  = Pt(0)


# ─── Title ───────────────────────────────────────────────────────────────────
h1('Wildfire and Long-Run Mental Health:\n'
   'Evidence from Matched Difference-in-Differences')
blank()
h1('[Author Name]  |  [Institution]  |  June 2026')
blank()
h1('References')
blank()

REFS = [
    ('Bergé, Laurent (2022). fixest: Fast Fixed Effects Estimation. '
     'R package version 0.10.4. https://lrberge.github.io/fixest/.'),

    ('Callaway, Brantly, and Pedro H. C. Sant’Anna (2021). '
     '“Difference-in-Differences with Multiple Time Periods.” '
     'Journal of Econometrics 225(2): 200–230. '
     'https://doi.org/10.1016/j.jeconom.2020.12.001.'),

    ('Currie, Janet, and Soodeh Saberian (2025). '
     '“Wildfire, Smoke and Mental Health in Canada.” '
     'NBER Working Paper No. 33912. https://www.nber.org/papers/w33912.'),

    ('Du, Ran, Ke Liu, Dangru Zhao, and Qiyun Fang (2024). '
     '“Climate Disaster and Cognitive Ability: Evidence From Wildfire.” '
     'International Journal of Public Health 69: 1607128. '
     'https://doi.org/10.3389/ijph.2024.1607128.'),

    ('Greenberg, Paul E., Andree-Anne Fournier, Tammy Sisitsky, Mark Simes, '
     'Richard Berman, Sarah H. Koenigsberg, and Ronald C. Kessler (2021). '
     '“The Economic Burden of Adults with Major Depressive Disorder '
     'in the United States (2010 and 2018).” '
     'PharmacoEconomics 39(6): 653–665. '
     'https://doi.org/10.1007/s40273-021-01019-4.'),

    ('Imbens, Guido W., and Michal Kolesár (2016). '
     '“Robust Standard Errors in Small Samples: Some Practical Advice.” '
     'Review of Economics and Statistics 98(4): 701–712. '
     'https://doi.org/10.1162/REST_a_00552.'),

    ('Jung, Youn Soo, Mary M. Johnson, Marshall Burke, Sara Heft-Neal, '
     'Mathieu Bondy, Sayantani Chinthrajah, Ashley C. Cullen, Kim Nelson, '
     'Nina Dresser, and Kari Nadeau (2025). '
     '“Fine Particulate Matter From 2020 California Wildfires and '
     'Mental Health–Related Emergency Department Visits.” '
     'JAMA Network Open 8(4): e253326. '
     'https://doi.org/10.1001/jamanetworkopen.2025.3326.'),

    ('Merdjanoff, Alexis A., Kate Burrows, and Kathleen Lynch (2026). '
     '“Projecting the Mental Health Impacts of the 2025 LA Wildfires: '
     'Lessons from Hurricane Katrina.” '
     'Environmental Research Letters 21(11): 111004. '
     'https://doi.org/10.1088/1748-9326/ae6d16.'),

    ('Pustejovsky, James E. (2022). clubSandwich: Cluster-Robust (Sandwich) '
     'Variance Estimators with Small-Sample Corrections. '
     'R package version 0.5.8. '
     'https://CRAN.R-project.org/package=clubSandwich.'),

    ('Wettstein, Zachary S., and Ambarish Vaidyanathan (2024). '
     '“Psychotropic Medication Prescriptions and Large California Wildfires.” '
     'JAMA Network Open 7(2): e2356466. '
     'https://doi.org/10.1001/jamanetworkopen.2023.56466.'),

    ('Ye, Xinyue, Yuning Ye, Xiao Huang, and Tracy Onega (2026). '
     '“Wildfires and Public Health: A Comprehensive Review of '
     'Human-Centric Studies.” '
     'GeoHealth 10(2): e2025GH001534. '
     'https://doi.org/10.1029/2025GH001534.'),
]

for entry in REFS:
    ref(entry)

out = r'C:\Users\chenyon\Research Paper 2026(1)\references.docx'
doc.save(out)
print('Saved:', out)
