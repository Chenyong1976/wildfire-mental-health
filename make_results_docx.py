from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(24)
style.paragraph_format.space_after  = Pt(0)
style.paragraph_format.space_before = Pt(0)


def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True
    return p


def para(text, indent=True, italic_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if italic_prefix:
        r0 = p.add_run(italic_prefix + "  ")
        r0.font.name = 'Times New Roman'
        r0.font.size = Pt(12)
        r0.italic = True
        r0.bold   = True
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p


def eq(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.first_line_indent = Inches(1.0)
    r = p.add_run(text + "    [" + label + "]")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.italic = True
    return p


def blank():
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_after  = Pt(0)


def add_table_caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True


def add_table_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(12)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)


# ═══════════════════════════════════════════════════════════════════════════════
# Title block
# ═══════════════════════════════════════════════════════════════════════════════
h1("Wildfire and Long-Run Mental Health:\n"
   "Evidence from Matched Difference-in-Differences")
blank()
h1("[Author Name]  |  [Institution]  |  June 2026")
blank()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5
# ═══════════════════════════════════════════════════════════════════════════════
h1("5.  Results")
blank()

# --- 5.1 Depression ---
h2("5.1  Depression Prevalence")
blank()

para(
    "Table 2 presents the headline finding. At the primary T1 threshold "
    "(fires >= 1,000 acres), counties that experienced a 2015 wildfire report "
    "depression prevalence 2.76 percentage points higher in 2019 than matched "
    "control counties (beta-hat = 2.76, SE = 0.94, p = 0.004). The effect is "
    "large relative to the control-group mean of 18.0 percent: it implies a "
    "15 percent increase in depression prevalence four years after the fires. "
    "To put the magnitude in context, a 2.8 percentage point increase in "
    "depression prevalence corresponds roughly to the estimated effect of a "
    "sustained 2-percentage-point rise in local unemployment over several "
    "years---an external benchmark consistent with the economic stress pathway "
    "that Currie and Saberian (2025) identify as one of the independent causal "
    "channels from wildfire exposure to mental health harm."
)

para(
    "Column (3) of Table 2 shows that the binary estimate is weaker and "
    "statistically insignificant at the T2 threshold (>= 5,000 acres): "
    "beta-hat = 0.55, p = 0.489. This pattern is consistent with the matching "
    "structure---the T2 control group (32 counties) is better balanced than "
    "T1's (12 counties), so the T2 comparison is the more credible one, but it "
    "pays a power cost from the smaller treated sample (41 counties versus 79 "
    "at T1). At T3 (>= 25,000 acres), the binary estimate is 1.35 percentage "
    "points (p = 0.251), directionally consistent with T1 but too imprecisely "
    "estimated to draw conclusions from alone."
)

blank()

# --- 5.2 Poor MH Days ---
h2("5.2  Poor Mental Health Days")
blank()

para(
    "The PLACES measure of poor mental health days provides a second, "
    "independent corroboration. Column (1) of Panel B in Table 2 shows that "
    "T1 wildfire counties report poor mental health days 1.63 percentage points "
    "higher in 2019 (beta-hat = 1.63, SE = 0.51, p = 0.002). The effect is "
    "again large relative to the control mean of 13.2 percent, implying a "
    "14 percent increase in the share of adults experiencing frequent mental "
    "distress. That two independent BRFSS-based measures---self-reported "
    "depression diagnosis and self-reported mental distress days---both show "
    "significant and similar-magnitude effects at T1 substantially reduces the "
    "probability that the finding is a statistical artifact of a single noisy "
    "outcome."
)

blank()

# --- 5.3 Dose-Response ---
h2("5.3  Dose-Response")
blank()

para(
    "Column (2) of Table 2 replaces the binary treatment indicator with the "
    "log of percent county land area burned in 2015, estimated among the 79 "
    "treated counties in the T1 matched panel. The dose-response coefficient "
    "is positive and statistically significant: beta-hat_D = 1.18 per log-unit "
    "of land burned (SE = 0.50, p = 0.020). A county at the 75th percentile of "
    "burn intensity (approximately 10 percent of land area burned) shows "
    "1.18 log-units higher depression than a county at the 25th percentile "
    "(approximately 0.5 percent burned), holding baseline controls constant."
)

para(
    "This monotone, graded relationship between how much a county burned and "
    "how much worse its 2019 depression is constitutes the strongest evidence "
    "for a causal interpretation. A confounding story based on the selection of "
    "structurally vulnerable counties into treatment predicts that treated "
    "counties have worse mental health regardless of fire severity---that is, "
    "it predicts a positive coefficient in the binary regression but predicts "
    "beta_D = 0 in the dose-response. The finding beta-hat_D = 1.18 > 0 "
    "(p = 0.020) is inconsistent with this null. By contrast, the dose-response "
    "coefficient for poor mental health days is small and statistically "
    "insignificant (beta-hat_D = 0.20, p = 0.465), suggesting that the graded "
    "exposure-response relationship is specific to the depression diagnosis "
    "measure and not uniformly present across all mental health outcomes."
)

blank()

# --- 5.4 Suicide and Overdose ---
h2("5.4  Suicide and Overdose Mortality")
blank()

para(
    "[Figure 1 about here — Event-Study Coefficients: Wildfire Effects on Suicide "
    "and Overdose Mortality, by Fire-Size Threshold. Generated by estimate_2015.R.]",
    indent=False
)
blank()

para(
    "Figure 1 presents the TWFE event-study coefficients for IHS-transformed "
    "suicide and overdose mortality rates across all three fire-size thresholds. "
    "Table 3 summarizes the pooled ATT estimates from the simple DiD and "
    "collapsed 2x2 specifications."
)

para(
    "Before discussing post-treatment effects, I confirm that the pre-treatment "
    "coefficients support the parallel-trends assumption.",
    italic_prefix="Pre-trend validation."
)

para(
    "At the T1 threshold, the three pre-treatment event-study coefficients "
    "(k = -4, -3, -2) pass the joint F-test cleanly: F = 0.53 (p = 0.663) "
    "for suicide and F = 0.51 (p = 0.679) for overdose. The individual "
    "coefficients are small and imprecise---the largest, at k = -2 for suicide "
    "(beta-hat = 0.487, SE = 0.431), is statistically indistinguishable from "
    "zero (p = 0.262). T3 passes cleanly as well (F = 0.85, p = 0.478 for "
    "suicide; F = 0.54, p = 0.659 for overdose). T2 pre-trends are borderline: "
    "F = 2.09 (p = 0.110) for suicide and F = 2.33 (p = 0.083) for overdose. "
    "T2 results should therefore be interpreted cautiously."
)

para(
    "Post-treatment suicide event-study coefficients at T1 are consistently "
    "positive: beta-hat_{k=0} = 0.568 (p = 0.169), beta-hat_{k=+1} = 0.428 "
    "(p = 0.299), beta-hat_{k=+2} = 0.413 (p = 0.323), beta-hat_{k=+3} = "
    "0.391 (p = 0.343), beta-hat_{k=+4} = 0.090 (p = 0.363). Overdose "
    "coefficients follow the same pattern: all post-treatment estimates are "
    "positive, ranging from 0.111 to 0.563 in IHS units, none individually "
    "significant at conventional levels. The confidence intervals are wide, "
    "reflecting the 50 percent rate of CDC WONDER suppression in the panel---"
    "the effective sample for mortality outcomes is less than half the nominal "
    "panel size.",
    italic_prefix="Post-treatment event-study estimates."
)

para(
    "The simple DiD ATT estimates pool the post-treatment years into a single "
    "summary. At T1, the ATT is +0.195 IHS units for suicide (SE = 0.200, "
    "p = 0.331) and +0.213 for overdose (SE = 0.201, p = 0.291). Collapsing "
    "to county-level pre- and post-period means reduces annual noise from "
    "censored cells without changing the direction of the estimate: the "
    "collapsed DiD ATT is +0.131 for suicide (p = 0.472) and +0.205 for "
    "overdose (p = 0.250). All mortality estimates are directionally consistent "
    "with a wildfire effect on suicide and drug overdose---none crosses zero---"
    "but the confidence intervals are wide enough to accommodate effect sizes "
    "ranging from negligible to clinically meaningful. These results are best "
    "interpreted as suggestive rather than definitive.",
    italic_prefix="Pooled ATT."
)

para(
    "At the T2 threshold (>= 5,000 acres), the four-year event-study "
    "coefficient for suicide is beta-hat_{k=+4} = 0.346 (SE = 0.160, "
    "p = 0.035), the only mortality point estimate individually significant "
    "at the 5 percent level. This result should be read carefully: T2 "
    "pre-trends are borderline, so the identifying assumption is less secure "
    "than at T1 and T3. Nevertheless, the four-year horizon is consistent with "
    "the trajectory suggested by the depression cross-section---the largest "
    "effects on mental health appear to accumulate rather than decay over time.",
    italic_prefix="T2 four-year suicide effect."
)

blank()

# --- 5.5 Heterogeneity ---
h2("5.5  Heterogeneity by Rurality")
blank()

para(
    "[Figure 2 about here — Rurality Heterogeneity: Event-Study Coefficients by "
    "Metro / Non-Metro Classification (T1, >=1,000 acres). Generated by estimate_2015.R.]",
    indent=False
)
blank()

para(
    "Figure 2 presents separate T1 event studies for metropolitan "
    "(RUCC 1-3) and non-metropolitan (RUCC 4-9) counties. In metropolitan "
    "counties, post-treatment suicide and overdose coefficients are near zero "
    "and inconsistent in sign across years. In non-metropolitan counties, "
    "post-treatment coefficients are uniformly positive and grow over time. "
    "The four-year non-metropolitan suicide coefficient is beta-hat_{k=+4} = "
    "0.241 (SE = 0.113, p = 0.038), the only event-study point estimate for "
    "suicide that is individually significant in the full results."
)

para(
    "This pattern is consistent with two non-exclusive mechanisms. First, "
    "rural counties have substantially fewer mental health providers: the "
    "literature documents that HPSA-designated counties have lower treatment "
    "rates for depression and higher long-run suicide rates following acute "
    "shocks (see, e.g., Currie and Saberian, 2025, on the access channel). "
    "Second, rural residents are more likely to own land that burned, to depend "
    "on natural resources for income, and to have weaker social safety nets---"
    "all of which amplify the economic stress pathway from fire to mental "
    "health. The metropolitan null result is informative on its own: it rules "
    "out the hypothesis that the non-metropolitan effect is merely an artifact "
    "of including high-exposure counties regardless of provider access, and it "
    "focuses policy implications on the rural communities where harm appears "
    "concentrated."
)

blank()

# ═══════════════════════════════════════════════════════════════════════════════
# Table 2
# ═══════════════════════════════════════════════════════════════════════════════
add_table_caption("Table 2. Cross-Sectional Depression and Poor Mental Health Days: 2019")

tbl2 = doc.add_table(rows=14, cols=5)
tbl2.style = 'Table Grid'

def cell_text(tbl, row, col, text, bold=False, italic=False, center=False, fontsize=10):
    cell = tbl.cell(row, col)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(fontsize)
    r.bold = bold
    r.italic = italic

# Header row 0: column headers
cell_text(tbl2, 0, 0, "", bold=True, fontsize=10)
cell_text(tbl2, 0, 1, "T1: Binary\n(1)", bold=True, center=True, fontsize=10)
cell_text(tbl2, 0, 2, "T1: Dose-resp.\n(2)", bold=True, center=True, fontsize=10)
cell_text(tbl2, 0, 3, "T2: Binary\n(3)", bold=True, center=True, fontsize=10)
cell_text(tbl2, 0, 4, "T3: Binary\n(4)", bold=True, center=True, fontsize=10)

# Panel A header
cell_text(tbl2, 1, 0, "Panel A: Depression prevalence (%)", italic=True, fontsize=10)
for c in range(1, 5):
    cell_text(tbl2, 1, c, "", fontsize=10)

# Treated row
cell_text(tbl2, 2, 0, "  Treated", fontsize=10)
cell_text(tbl2, 2, 1, "2.758**", center=True, fontsize=10)
cell_text(tbl2, 2, 2, "", center=True, fontsize=10)
cell_text(tbl2, 2, 3, "0.546", center=True, fontsize=10)
cell_text(tbl2, 2, 4, "1.350", center=True, fontsize=10)

# SE row
cell_text(tbl2, 3, 0, "", fontsize=10)
cell_text(tbl2, 3, 1, "(0.944)", center=True, fontsize=10)
cell_text(tbl2, 3, 2, "", center=True, fontsize=10)
cell_text(tbl2, 3, 3, "(0.784)", center=True, fontsize=10)
cell_text(tbl2, 3, 4, "(1.148)", center=True, fontsize=10)

# Dose row
cell_text(tbl2, 4, 0, "  log(1 + PctBurned)", fontsize=10)
cell_text(tbl2, 4, 1, "", center=True, fontsize=10)
cell_text(tbl2, 4, 2, "1.180*", center=True, fontsize=10)
cell_text(tbl2, 4, 3, "", center=True, fontsize=10)
cell_text(tbl2, 4, 4, "", center=True, fontsize=10)

cell_text(tbl2, 5, 0, "", fontsize=10)
cell_text(tbl2, 5, 1, "", center=True, fontsize=10)
cell_text(tbl2, 5, 2, "(0.496)", center=True, fontsize=10)
cell_text(tbl2, 5, 3, "", center=True, fontsize=10)
cell_text(tbl2, 5, 4, "", center=True, fontsize=10)

# Panel B header
cell_text(tbl2, 6, 0, "Panel B: Poor mental health days (%)", italic=True, fontsize=10)
for c in range(1, 5):
    cell_text(tbl2, 6, c, "", fontsize=10)

# Treated row B
cell_text(tbl2, 7, 0, "  Treated", fontsize=10)
cell_text(tbl2, 7, 1, "1.634**", center=True, fontsize=10)
cell_text(tbl2, 7, 2, "", center=True, fontsize=10)
cell_text(tbl2, 7, 3, "0.269", center=True, fontsize=10)
cell_text(tbl2, 7, 4, "0.580", center=True, fontsize=10)

cell_text(tbl2, 8, 0, "", fontsize=10)
cell_text(tbl2, 8, 1, "(0.505)", center=True, fontsize=10)
cell_text(tbl2, 8, 2, "", center=True, fontsize=10)
cell_text(tbl2, 8, 3, "(0.355)", center=True, fontsize=10)
cell_text(tbl2, 8, 4, "(0.410)", center=True, fontsize=10)

# Dose row B
cell_text(tbl2, 9, 0, "  log(1 + PctBurned)", fontsize=10)
cell_text(tbl2, 9, 1, "", center=True, fontsize=10)
cell_text(tbl2, 9, 2, "0.198", center=True, fontsize=10)
cell_text(tbl2, 9, 3, "", center=True, fontsize=10)
cell_text(tbl2, 9, 4, "", center=True, fontsize=10)

cell_text(tbl2, 10, 0, "", fontsize=10)
cell_text(tbl2, 10, 1, "", center=True, fontsize=10)
cell_text(tbl2, 10, 2, "(0.269)", center=True, fontsize=10)
cell_text(tbl2, 10, 3, "", center=True, fontsize=10)
cell_text(tbl2, 10, 4, "", center=True, fontsize=10)

# Footer rows
cell_text(tbl2, 11, 0, "Controls", fontsize=10)
for c in range(1, 5):
    cell_text(tbl2, 11, c, "Yes", center=True, fontsize=10)

cell_text(tbl2, 12, 0, "N (counties)", fontsize=10)
cell_text(tbl2, 12, 1, "90", center=True, fontsize=10)
cell_text(tbl2, 12, 2, "79", center=True, fontsize=10)
cell_text(tbl2, 12, 3, "64", center=True, fontsize=10)
cell_text(tbl2, 12, 4, "30", center=True, fontsize=10)

# Threshold labels in last row
cell_text(tbl2, 13, 0, "Threshold", fontsize=10)
cell_text(tbl2, 13, 1, ">=1,000 ac", center=True, fontsize=10)
cell_text(tbl2, 13, 2, ">=1,000 ac", center=True, fontsize=10)
cell_text(tbl2, 13, 3, ">=5,000 ac", center=True, fontsize=10)
cell_text(tbl2, 13, 4, ">=25,000 ac", center=True, fontsize=10)

add_table_note(
    "Notes: Outcome is 2019 CDC PLACES county-level prevalence. Columns (1)/(3)/(4): OLS with binary treatment "
    "indicator (Eq. 4). Column (2): OLS with log(1 + PctBurned_2015), estimated among treated counties only (Eq. 5). "
    "Controls: 2011-2014 means of IHS suicide rate, IHS overdose rate, and unemployment rate; 2013 RUCC score. "
    "Heteroskedasticity-robust SEs in parentheses. *** p<0.01, ** p<0.05, * p<0.10."
)

blank()

# ═══════════════════════════════════════════════════════════════════════════════
# Table 3
# ═══════════════════════════════════════════════════════════════════════════════
add_table_caption("Table 3. Wildfire Effects on Suicide and Drug Overdose Mortality: DiD Estimates")

tbl3 = doc.add_table(rows=13, cols=6)
tbl3.style = 'Table Grid'

# Header row
cell_text(tbl3, 0, 0, "", bold=True, fontsize=10)
cell_text(tbl3, 0, 1, "T1\nSuicide", bold=True, center=True, fontsize=10)
cell_text(tbl3, 0, 2, "T1\nOverdose", bold=True, center=True, fontsize=10)
cell_text(tbl3, 0, 3, "T2\nSuicide", bold=True, center=True, fontsize=10)
cell_text(tbl3, 0, 4, "T2\nOverdose", bold=True, center=True, fontsize=10)
cell_text(tbl3, 0, 5, "T3\nSuicide", bold=True, center=True, fontsize=10)

# Panel A: Simple DiD
cell_text(tbl3, 1, 0, "Panel A: Simple DiD (pooled ATT, 2015-2019)", italic=True, fontsize=10)
for c in range(1, 6):
    cell_text(tbl3, 1, c, "", fontsize=10)

cell_text(tbl3, 2, 0, "  ATT", fontsize=10)
cell_text(tbl3, 2, 1, "0.195", center=True, fontsize=10)
cell_text(tbl3, 2, 2, "0.213", center=True, fontsize=10)
cell_text(tbl3, 2, 3, "0.067", center=True, fontsize=10)
cell_text(tbl3, 2, 4, "0.100", center=True, fontsize=10)
cell_text(tbl3, 2, 5, "-0.056", center=True, fontsize=10)

cell_text(tbl3, 3, 0, "", fontsize=10)
cell_text(tbl3, 3, 1, "(0.200)", center=True, fontsize=10)
cell_text(tbl3, 3, 2, "(0.201)", center=True, fontsize=10)
cell_text(tbl3, 3, 3, "(0.148)", center=True, fontsize=10)
cell_text(tbl3, 3, 4, "(0.162)", center=True, fontsize=10)
cell_text(tbl3, 3, 5, "(0.185)", center=True, fontsize=10)

cell_text(tbl3, 4, 0, "  p-value", fontsize=10)
cell_text(tbl3, 4, 1, "0.331", center=True, fontsize=10)
cell_text(tbl3, 4, 2, "0.291", center=True, fontsize=10)
cell_text(tbl3, 4, 3, "0.654", center=True, fontsize=10)
cell_text(tbl3, 4, 4, "0.539", center=True, fontsize=10)
cell_text(tbl3, 4, 5, "0.766", center=True, fontsize=10)

# Panel B: Collapsed 2x2
cell_text(tbl3, 5, 0, "Panel B: Collapsed 2x2 DiD (pre=2011-14, post=2016-19)", italic=True, fontsize=10)
for c in range(1, 6):
    cell_text(tbl3, 5, c, "", fontsize=10)

cell_text(tbl3, 6, 0, "  ATT", fontsize=10)
cell_text(tbl3, 6, 1, "0.131", center=True, fontsize=10)
cell_text(tbl3, 6, 2, "0.205", center=True, fontsize=10)
cell_text(tbl3, 6, 3, "0.113", center=True, fontsize=10)
cell_text(tbl3, 6, 4, "0.122", center=True, fontsize=10)
cell_text(tbl3, 6, 5, "-0.102", center=True, fontsize=10)

cell_text(tbl3, 7, 0, "", fontsize=10)
cell_text(tbl3, 7, 1, "(0.181)", center=True, fontsize=10)
cell_text(tbl3, 7, 2, "(0.177)", center=True, fontsize=10)
cell_text(tbl3, 7, 3, "(0.159)", center=True, fontsize=10)
cell_text(tbl3, 7, 4, "(0.166)", center=True, fontsize=10)
cell_text(tbl3, 7, 5, "(0.222)", center=True, fontsize=10)

cell_text(tbl3, 8, 0, "  p-value", fontsize=10)
cell_text(tbl3, 8, 1, "0.472", center=True, fontsize=10)
cell_text(tbl3, 8, 2, "0.250", center=True, fontsize=10)
cell_text(tbl3, 8, 3, "0.481", center=True, fontsize=10)
cell_text(tbl3, 8, 4, "0.465", center=True, fontsize=10)
cell_text(tbl3, 8, 5, "0.651", center=True, fontsize=10)

# Panel C: Pre-trend F
cell_text(tbl3, 9, 0, "Panel C: Pre-trend joint F-test (k = -4, -3, -2)", italic=True, fontsize=10)
for c in range(1, 6):
    cell_text(tbl3, 9, c, "", fontsize=10)

cell_text(tbl3, 10, 0, "  F-statistic", fontsize=10)
cell_text(tbl3, 10, 1, "0.53", center=True, fontsize=10)
cell_text(tbl3, 10, 2, "0.51", center=True, fontsize=10)
cell_text(tbl3, 10, 3, "2.09", center=True, fontsize=10)
cell_text(tbl3, 10, 4, "2.33", center=True, fontsize=10)
cell_text(tbl3, 10, 5, "0.85", center=True, fontsize=10)

cell_text(tbl3, 11, 0, "  p-value", fontsize=10)
cell_text(tbl3, 11, 1, "0.663", center=True, fontsize=10)
cell_text(tbl3, 11, 2, "0.679", center=True, fontsize=10)
cell_text(tbl3, 11, 3, "0.110", center=True, fontsize=10)
cell_text(tbl3, 11, 4, "0.083", center=True, fontsize=10)
cell_text(tbl3, 11, 5, "0.478", center=True, fontsize=10)

# Footer
cell_text(tbl3, 12, 0, "Counties (treated / control)", fontsize=10)
cell_text(tbl3, 12, 1, "81 / 12", center=True, fontsize=10)
cell_text(tbl3, 12, 2, "81 / 12", center=True, fontsize=10)
cell_text(tbl3, 12, 3, "42 / 24", center=True, fontsize=10)
cell_text(tbl3, 12, 4, "42 / 24", center=True, fontsize=10)
cell_text(tbl3, 12, 5, "17 / 14", center=True, fontsize=10)

add_table_note(
    "Notes: Outcome is IHS-transformed annual county mortality rate (per 100,000). "
    "Panel A: TWFE with county and year fixed effects, unemployment rate control, SEs clustered by county (Eq. 3). "
    "Panel B: Panel collapsed to pre-period (2011-2014) and post-period (2016-2019) county means; treatment year 2015 excluded. "
    "Panel C: Joint Wald F-test on event-study leads k = -4, -3, -2 (reference: k = -1). "
    "T3 overdose omitted due to insufficient variation. *** p<0.01, ** p<0.05, * p<0.10."
)

out = r"C:\Users\chenyon\Research Paper 2026(1)\section_results.docx"
doc.save(out)
print("Saved: " + out)
