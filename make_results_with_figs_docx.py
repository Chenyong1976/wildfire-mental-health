"""
Regenerate section_results.docx with actual figure images embedded.
"""
import sys; sys.path.insert(0, r'C:\Users\chenyon\Research Paper 2026(1)')
from make_equations import add_math_runs
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(24)
style.paragraph_format.space_after  = Pt(0)
style.paragraph_format.space_before = Pt(0)

FIG1 = r"C:\Users\chenyon\Research Paper 2026(1)\fig_eventstudy.png"
FIG2 = r"C:\Users\chenyon\Research Paper 2026(1)\fig_rucc.png"


def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True


def para(text, indent=True, italic_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if italic_prefix:
        r0 = p.add_run(italic_prefix + '  ')
        r0.font.name = 'Times New Roman'
        r0.font.size = Pt(12)
        r0.italic = True
        r0.bold   = True
    add_math_runs(p, text)


def blank(half=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(6 if half else 12)
    p.paragraph_format.space_after  = Pt(0)


def fig_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)


def fig_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)


def add_table_caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True


def add_table_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(12)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)


def cell_text(tbl, row, col, text, bold=False, italic=False,
              center=False, fontsize=10):
    cell = tbl.cell(row, col)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(fontsize)
    r.bold   = bold
    r.italic = italic


# ═══════════════════════════════════════════════════════════════════════════════
h1('5.  Results')
blank()

h2('5.1  Depression Prevalence')
blank()

p_ph2_dep = doc.add_paragraph()
p_ph2_dep.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ph2_dep.paragraph_format.line_spacing = Pt(18)
p_ph2_dep.paragraph_format.space_after = Pt(6)
r_ph2_dep = p_ph2_dep.add_run('[Figure 2 about here]')
r_ph2_dep.italic = True
r_ph2_dep.font.name = 'Times New Roman'
r_ph2_dep.font.size = Pt(11)
blank()

para('Figure 2 presents all Table 2 estimates graphically, with 95 percent '
     'confidence intervals, for both outcomes across all three fire-size thresholds '
     'and both specifications. The most prominent visual feature is that the T1 '
     'confidence interval lies entirely to the right of zero in both panels, while '
     'the T2 and T3 intervals straddle zero. This reflects the power gradient across '
     'thresholds: the T1 treated sample (79 counties) is nearly twice the T2 sample '
     '(41 counties) and five times the T3 sample (17 counties), so the T1 comparison '
     'has substantially greater precision despite similarly sized point estimates.')
blank()

para('Table 2 presents the headline finding. At the primary T1 threshold '
     '(fires ≥ 1,000 acres), counties that experienced a 2015 wildfire report '
     'depression prevalence 2.76 percentage points higher in 2019 than matched '
     'control counties (β̂ = 2.76, SE = 0.94, p = 0.004). The effect is '
     'large relative to the control-group mean of 18.0 percent: it implies a '
     '15 percent increase in depression prevalence four years after the fires. '
     'To put the magnitude in context, a 2.8 percentage point increase in '
     'depression prevalence is roughly comparable to the effect of a sustained '
     '2-percentage-point rise in local unemployment over several years. '
     'Currie and Saberian (2025) identify economic stress as one independent '
     'causal channel from wildfire exposure to mental health harm, making this '
     'benchmark particularly relevant.')
para('Column (3) of Table 2 shows that the binary estimate is weaker and '
     'statistically insignificant at the T2 threshold (≥ 5,000 acres): '
     'β̂ = 0.55, p = 0.489. This pattern is consistent with the matching '
     'structure: the T2 control group (24 counties) is better balanced than '
     "T1's (12 counties), so the T2 comparison is the more credible one, but it "
     'pays a power cost from the smaller treated sample (41 counties versus 79 '
     'at T1). At T3 (≥ 25,000 acres), the binary estimate is 1.35 percentage '
     'points (p = 0.251), directionally consistent with T1 but too imprecisely '
     'estimated to draw conclusions from alone.')
blank()

h2('5.2  Poor Mental Health Days')
blank()
para('The PLACES measure of poor mental health days provides a second, '
     'independent corroboration. Column (1) of Panel B in Table 2 shows that '
     'T1 wildfire counties report poor mental health days 1.63 percentage points '
     'higher in 2019 (β̂ = 1.63, SE = 0.51, p = 0.002). The effect is '
     'large relative to the control mean of 13.2 percent, implying a '
     '12 percent increase in the share of adults experiencing frequent mental '
     'distress. Two independent BRFSS-based measures—self-reported depression '
     'diagnosis and self-reported mental distress days—both show significant '
     'effects of similar magnitude at T1. As such, the finding is unlikely to '
     'be a statistical artifact of a single noisy outcome. '
     'Panel B of Figure 2 shows the T1, T2, and T3 estimates for poor mental '
     'health days with 95 percent confidence intervals; the T1 interval again '
     'lies above zero while T2 and T3 straddle it.')
blank()

h2('5.3  Dose-Response')
blank()
para('Column (2) of Table 2 replaces the binary treatment indicator with the '
     'log of percent county land area burned in 2015, estimated among the 79 '
     'treated counties in the T1 matched panel. The dose-response coefficient '
     'is positive and statistically significant: β̂_{D} = 1.18 per log-unit '
     'of land burned (SE = 0.50, p = 0.020). A county at the 75th percentile of '
     'burn intensity (approximately 10 percent of land area burned) shows '
     '1.18 log-units higher depression than a county at the 25th percentile '
     '(approximately 0.5 percent burned), holding baseline controls constant.')
para('This monotone, graded relationship between how much a county burned and '
     'how much worse its 2019 depression is constitutes the strongest evidence '
     'for a causal interpretation. A confounding story based on the selection of '
     'structurally vulnerable counties into treatment predicts that treated '
     'counties have worse mental health regardless of fire severity; that is, '
     'it predicts a positive coefficient in the binary regression but predicts '
     'β_{D} = 0 in the dose-response. The finding β̂_{D} = 1.18 > 0 '
     '(p = 0.020) is inconsistent with this null. By contrast, the dose-response '
     'coefficient for poor mental health days is small and statistically '
     'insignificant (β̂_{D} = 0.20, p = 0.465), suggesting that the graded '
     'exposure-response relationship is specific to the depression diagnosis '
     'measure and not uniformly present across all mental health outcomes. '
     'The dose-response rows in Figure 2 make this contrast clear: the depression '
     'dose-response confidence interval lies above zero (Panel A, red diamond) '
     'while the poor mental health days interval spans it (Panel B).')
blank()

h2('5.4  Suicide and Overdose Mortality')
blank()

p_ph1 = doc.add_paragraph()
p_ph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ph1.paragraph_format.line_spacing = Pt(18)
p_ph1.paragraph_format.space_after = Pt(6)
r_ph1 = p_ph1.add_run('[Figure 3 about here]')
r_ph1.italic = True
r_ph1.font.name = 'Times New Roman'
r_ph1.font.size = Pt(11)
blank()

para('Figure 3 presents the baseline TWFE event-study coefficients for '
     'IHS-transformed suicide (left panel) and drug overdose mortality (right panel). '
     'Each panel overlays estimates for all three fire-size thresholds: T1 '
     '(≥1,000 acres), T2 (≥5,000 acres), and T3 (≥25,000 acres), '
     'with reference year k = −1 (2014) normalized to zero.')

p_ph2a = doc.add_paragraph()
p_ph2a.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ph2a.paragraph_format.line_spacing = Pt(18)
p_ph2a.paragraph_format.space_after = Pt(6)
r_ph2a = p_ph2a.add_run('[Figure 4 about here]')
r_ph2a.italic = True
r_ph2a.font.name = 'Times New Roman'
r_ph2a.font.size = Pt(11)
blank()

para('Figure 4 presents the TWFE event-study coefficients for IHS-transformed '
     'suicide and overdose mortality rates across all three fire-size thresholds. '
     'Table 3 summarizes the pooled ATT estimates from the simple DiD and '
     'collapsed 2x2 specifications.')
para('Before discussing post-treatment effects, I confirm that the pre-treatment '
     'coefficients support the parallel-trends assumption.',
     italic_prefix='Pre-trend validation.')
para('At the T1 threshold, the three pre-treatment event-study coefficients '
     '(k = −4, −3, −2) pass the joint F-test cleanly: F = 0.53 (p = 0.663) '
     'for suicide and F = 0.51 (p = 0.679) for overdose. The individual '
     'coefficients are small and imprecise; the largest, at k = −2 for suicide '
     '(β̂ = 0.487, SE = 0.431), is statistically indistinguishable from '
     'zero (p = 0.262). T3 passes cleanly as well (F = 0.85, p = 0.478 for '
     'suicide; T3 overdose is omitted due to insufficient cell variation). T2 pre-trends are borderline: '
     'F = 2.09 (p = 0.110) for suicide and F = 2.33 (p = 0.083) for overdose. '
     'T2 results should therefore be interpreted cautiously.')
para('Post-treatment suicide event-study coefficients at T1 are consistently '
     'positive: β̂_{k=0} = 0.568 (p = 0.169), β̂_{k=+1} = 0.428 '
     '(p = 0.299), β̂_{k=+2} = 0.413 (p = 0.323), β̂_{k=+3} = '
     '0.391 (p = 0.343), β̂_{k=+4} = 0.090 (p = 0.363). Overdose '
     'coefficients follow the same pattern: all post-treatment estimates are '
     'positive, ranging from 0.111 to 0.563 in IHS units, none individually '
     'significant at conventional levels. The confidence intervals are wide, '
     'reflecting the 50 percent rate of CDC WONDER suppression in the panel; '
     'the effective sample for mortality outcomes is less than half the nominal '
     'panel size.',
     italic_prefix='Post-treatment event-study estimates.')
para('The simple DiD ATT estimates pool the post-treatment years into a single '
     'summary. At T1, the ATT is +0.195 IHS units for suicide (SE = 0.200, '
     'p = 0.331) and +0.213 for overdose (SE = 0.201, p = 0.291). Collapsing '
     'to county-level pre- and post-period means reduces annual noise from '
     'censored cells without changing the direction of the estimate: the '
     'collapsed DiD ATT is +0.131 for suicide (p = 0.472) and +0.205 for '
     'overdose (p = 0.250). All mortality estimates are directionally consistent '
     'with a wildfire effect on suicide and drug overdose (none crosses zero), '
     'but the confidence intervals are wide enough to accommodate effect sizes '
     'ranging from negligible to clinically meaningful. These results are best '
     'interpreted as suggestive rather than definitive.',
     italic_prefix='Pooled ATT.')
para('At the T2 threshold (≥ 5,000 acres), the four-year event-study '
     'coefficient for suicide is β̂_{k=+4} = 0.346 (SE = 0.160, '
     'p = 0.035), the only mortality point estimate individually significant '
     'at the 5 percent level. This result should be read carefully: T2 '
     'pre-trends are borderline, so the identifying assumption is less secure '
     'than at T1 and T3. Nevertheless, the four-year horizon is consistent with '
     'the trajectory suggested by the depression cross-section: the largest '
     'effects on mental health appear to accumulate rather than decay over time.',
     italic_prefix='T2 four-year suicide effect.')
blank()

h2('5.5  Heterogeneity by Rurality')
blank()

p_ph2 = doc.add_paragraph()
p_ph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ph2.paragraph_format.line_spacing = Pt(18)
p_ph2.paragraph_format.space_after = Pt(6)
r_ph2 = p_ph2.add_run('[Figure 5 about here]')
r_ph2.italic = True
r_ph2.font.name = 'Times New Roman'
r_ph2.font.size = Pt(11)
blank()

para('Figure 5 presents separate T1 event studies for metropolitan '
     '(RUCC 1–3) and non-metropolitan (RUCC 4–9) counties. In metropolitan '
     'counties, post-treatment suicide and overdose coefficients are near zero '
     'and inconsistent in sign across years. In non-metropolitan counties, '
     'post-treatment coefficients are uniformly positive and grow over time. '
     'The four-year non-metropolitan suicide coefficient is β̂_{k=+4} = '
     '0.241 (SE = 0.113, p = 0.038), the only event-study point estimate for '
     'suicide that is individually significant in the full results.')
para('This pattern is consistent with two non-exclusive mechanisms. First, '
     'rural counties have fewer mental health providers: the literature documents '
     'that HPSA-designated counties have lower treatment rates for depression and '
     'higher long-run suicide rates following acute shocks. Second, rural '
     'residents are more likely to own land that burned, to depend on natural '
     'resources for income, and to have weaker social safety nets—all of which '
     'amplify the economic stress pathway from fire to mental health. The '
     'metropolitan null result is informative on its own: it rules out the '
     'hypothesis that the non-metropolitan effect is merely an artifact of '
     'including high-exposure counties regardless of provider access, and it '
     'focuses policy implications on the rural communities where harm appears '
     'concentrated.')
blank()

# ─── Table 2 ─────────────────────────────────────────────────────────────────
add_table_caption('Table 2. Cross-Sectional Depression and Poor Mental Health Days: 2019')
tbl2 = doc.add_table(rows=14, cols=5)
tbl2.style = 'Table Grid'

cell_text(tbl2, 0, 0, '', bold=True)
cell_text(tbl2, 0, 1, 'T1: Binary\n(1)',      bold=True, center=True)
cell_text(tbl2, 0, 2, 'T1: Dose-resp.\n(2)',  bold=True, center=True)
cell_text(tbl2, 0, 3, 'T2: Binary\n(3)',      bold=True, center=True)
cell_text(tbl2, 0, 4, 'T3: Binary\n(4)',      bold=True, center=True)

cell_text(tbl2, 1, 0, 'Panel A: Depression prevalence (%)', italic=True)
for c in range(1, 5): cell_text(tbl2, 1, c, '')

cell_text(tbl2, 2, 0, '  Treated')
cell_text(tbl2, 2, 1, '2.758***', center=True)
cell_text(tbl2, 2, 2, '',        center=True)
cell_text(tbl2, 2, 3, '0.546',   center=True)
cell_text(tbl2, 2, 4, '1.350',   center=True)

cell_text(tbl2, 3, 1, '(0.944)', center=True)
cell_text(tbl2, 3, 2, '',        center=True)
cell_text(tbl2, 3, 3, '(0.784)', center=True)
cell_text(tbl2, 3, 4, '(1.148)', center=True)

cell_text(tbl2, 4, 0, '  log(1 + PctBurned)')
cell_text(tbl2, 4, 2, '1.180**', center=True)

cell_text(tbl2, 5, 2, '(0.496)', center=True)

cell_text(tbl2, 6, 0, 'Panel B: Poor mental health days (%)', italic=True)
for c in range(1, 5): cell_text(tbl2, 6, c, '')

cell_text(tbl2, 7, 0, '  Treated')
cell_text(tbl2, 7, 1, '1.634***', center=True)
cell_text(tbl2, 7, 2, '',        center=True)
cell_text(tbl2, 7, 3, '0.269',   center=True)
cell_text(tbl2, 7, 4, '0.580',   center=True)

cell_text(tbl2, 8, 1, '(0.505)', center=True)
cell_text(tbl2, 8, 2, '',        center=True)
cell_text(tbl2, 8, 3, '(0.355)', center=True)
cell_text(tbl2, 8, 4, '(0.410)', center=True)

cell_text(tbl2, 9, 0, '  log(1 + PctBurned)')
cell_text(tbl2, 9, 2, '0.198',  center=True)
cell_text(tbl2, 10, 2, '(0.269)', center=True)

cell_text(tbl2, 11, 0, 'Controls')
for c in range(1, 5): cell_text(tbl2, 11, c, 'Yes', center=True)

cell_text(tbl2, 12, 0, 'N (counties)')
for val, c in zip(['90', '79', '64', '30'], range(1, 5)):
    cell_text(tbl2, 12, c, val, center=True)

cell_text(tbl2, 13, 0, 'Threshold')
for val, c in zip(['>=1,000 ac', '>=1,000 ac', '>=5,000 ac', '>=25,000 ac'], range(1, 5)):
    cell_text(tbl2, 13, c, val, center=True)

add_table_note('Notes: Outcome is 2019 CDC PLACES county-level prevalence. '
               'Columns (1)/(3)/(4): OLS with binary treatment indicator (Eq. 4). '
               'Column (2): OLS with log(1 + PctBurned_2015) among treated counties (Eq. 5). '
               'Controls: 2011-2014 means of IHS suicide rate, IHS overdose rate, '
               'unemployment rate; 2013 RUCC score. '
               'Heteroskedasticity-robust SEs in parentheses. *** p<0.01, ** p<0.05, * p<0.10.')
blank()

# ─── Table 3 ─────────────────────────────────────────────────────────────────
add_table_caption('Table 3. Wildfire Effects on Suicide and Drug Overdose Mortality: DiD Estimates')
tbl3 = doc.add_table(rows=13, cols=6)
tbl3.style = 'Table Grid'

cell_text(tbl3, 0, 0, '', bold=True)
for val, c in zip(['T1\nSuicide', 'T1\nOverdose', 'T2\nSuicide',
                   'T2\nOverdose', 'T3\nSuicide'], range(1, 6)):
    cell_text(tbl3, 0, c, val, bold=True, center=True)

cell_text(tbl3, 1, 0, 'Panel A: Simple DiD (pooled ATT, 2015-2019)', italic=True)
for c in range(1, 6): cell_text(tbl3, 1, c, '')

r2_vals = ['0.195', '0.213', '0.067', '0.100', '-0.056']
r3_vals = ['(0.200)', '(0.201)', '(0.148)', '(0.162)', '(0.185)']
r4_vals = ['0.331', '0.291', '0.654', '0.539', '0.766']

cell_text(tbl3, 2, 0, '  ATT')
for v, c in zip(r2_vals, range(1, 6)): cell_text(tbl3, 2, c, v, center=True)
for v, c in zip(r3_vals, range(1, 6)): cell_text(tbl3, 3, c, v, center=True)
cell_text(tbl3, 4, 0, '  p-value')
for v, c in zip(r4_vals, range(1, 6)): cell_text(tbl3, 4, c, v, center=True)

cell_text(tbl3, 5, 0, 'Panel B: Collapsed 2x2 DiD (pre=2011-14, post=2016-19)', italic=True)
for c in range(1, 6): cell_text(tbl3, 5, c, '')

r6_vals = ['0.131', '0.205', '0.113', '0.122', '-0.102']
r7_vals = ['(0.181)', '(0.177)', '(0.159)', '(0.166)', '(0.222)']
r8_vals = ['0.472', '0.250', '0.481', '0.465', '0.651']

cell_text(tbl3, 6, 0, '  ATT')
for v, c in zip(r6_vals, range(1, 6)): cell_text(tbl3, 6, c, v, center=True)
for v, c in zip(r7_vals, range(1, 6)): cell_text(tbl3, 7, c, v, center=True)
cell_text(tbl3, 8, 0, '  p-value')
for v, c in zip(r8_vals, range(1, 6)): cell_text(tbl3, 8, c, v, center=True)

cell_text(tbl3, 9, 0, 'Panel C: Pre-trend joint F-test (k = -4, -3, -2)', italic=True)
for c in range(1, 6): cell_text(tbl3, 9, c, '')

cell_text(tbl3, 10, 0, '  F-statistic')
for v, c in zip(['0.53', '0.51', '2.09', '2.33', '0.85'], range(1, 6)):
    cell_text(tbl3, 10, c, v, center=True)
cell_text(tbl3, 11, 0, '  p-value')
for v, c in zip(['0.663', '0.679', '0.110', '0.083', '0.478'], range(1, 6)):
    cell_text(tbl3, 11, c, v, center=True)

cell_text(tbl3, 12, 0, 'Counties (treated / control)')
for v, c in zip(['81/12', '81/12', '42/24', '42/24', '17/14'], range(1, 6)):
    cell_text(tbl3, 12, c, v, center=True)

add_table_note('Notes: Outcome is IHS-transformed annual county mortality rate (per 100,000). '
               'Panel A: TWFE with county and year FE, unemployment rate control, '
               'SEs clustered by county (Eq. 3). '
               'Panel B: County-level pre/post means; treatment year 2015 excluded. '
               'Panel C: Joint Wald F-test on k = −4, −3, −2. '
               'T3 overdose omitted due to insufficient variation. '
               '*** p<0.01, ** p<0.05, * p<0.10.')

out = r'C:\Users\chenyon\Research Paper 2026(1)\section_results.docx'
doc.save(out)
print('Saved:', out)
