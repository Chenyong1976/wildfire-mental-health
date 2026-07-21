import sys; sys.path.insert(0, r'C:\Users\chenyon\Research Paper 2026(1)')
from make_equations import add_math_runs
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(24)
style.paragraph_format.space_after  = Pt(0)
style.paragraph_format.space_before = Pt(0)


def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True


def para(text, indent=True, italic_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if italic_prefix:
        r0 = p.add_run(italic_prefix + "  ")
        r0.font.name = 'Times New Roman'
        r0.font.size = Pt(12)
        r0.italic = True
        r0.bold   = True
    add_math_runs(p, text)


def blank():
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_after  = Pt(0)


def caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True


def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(12)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)


def cell_text(tbl, row, col, text, bold=False, italic=False, center=False, fs=10):
    cell = tbl.cell(row, col)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(fs)
    r.bold   = bold
    r.italic = italic


# ─── Section header ───────────────────────────────────────────────────────────
h1("6.  Robustness")
blank()

# ─── 6.1 ─────────────────────────────────────────────────────────────────────
h2("6.1  Alternative Fire-Size Thresholds")
blank()

para(
    "Section 5 reports results at three fire-size thresholds (T1 >= 1,000 acres, "
    "T2 >= 5,000 acres, T3 >= 25,000 acres). Here I summarize what the "
    "cross-threshold pattern implies for identification."
)
para(
    "The depression finding is specific to T1 in terms of statistical significance: "
    "at T1, the binary OLS estimate of +2.76 percentage points is significant at "
    "the 1 percent level; at T2 and T3, binary estimates are positive but "
    "insignificant (p = 0.489 and p = 0.251, respectively). This is not evidence "
    "against a causal effect; the attenuation reflects power loss as the treated "
    "sample shrinks from 79 to 41 to 17 counties, but it means the depression "
    "finding rests primarily on the T1 analysis. The dose-response specification, "
    "which compares intensity within the T1 treated sample, remains the strongest "
    "within-threshold causal test, as it is not subject to the between-threshold "
    "power argument."
)
para(
    "For suicide and overdose mortality, all three thresholds produce positive "
    "pooled ATTs with the exception of T3 suicide under the simple DiD "
    "(ATT = −0.056, p = 0.766), which is indistinguishable from zero at either "
    "sign. The T2 four-year suicide estimate (β̂_{k=+4} = 0.346, p = 0.035) "
    "is the only mortality estimate to achieve individual significance, and it "
    "should be interpreted cautiously given the borderline T2 pre-trend test. "
    "The overall pattern (positive depression effects concentrated at T1, "
    "mortality coefficients uniformly positive but imprecise) is consistent across "
    "all three thresholds and rules out a scenario in which results are driven by "
    "any single threshold choice."
)

blank()

# ─── 6.2 ─────────────────────────────────────────────────────────────────────
h2("6.2  Collapsed 2x2 DiD")
blank()

para(
    "Panel B of Table 3 reports the collapsed 2x2 DiD, in which each county is "
    "reduced to a pre-period mean (2011-2014) and a post-period mean (2016-2019) "
    "before estimating the simple DiD. Collapsing addresses two specific concerns "
    "about the simple DiD in a context with heavy cell suppression. First, annual "
    "suppressed cells introduce non-classical measurement error that may bias the "
    "TWFE estimate if suppression is correlated with treatment; collapsing to "
    "multi-year means averages out annual suppression noise. Second, with only 12 "
    "control counties at T1, year-level fluctuations in the control group may "
    "receive excessive weight; averaging to two observations per county reduces "
    "this influence."
)
para(
    "The collapsed estimates are directionally consistent with the simple DiD at "
    "all thresholds and for both outcomes. At T1, the collapsed ATT is +0.131 for "
    "suicide (p = 0.472) and +0.205 for overdose (p = 0.250), slightly smaller "
    "than the simple DiD ATTs (+0.195 and +0.213), reflecting the removal of the "
    "treatment year 2015 from the post-period mean. The qualitative conclusion is "
    "unchanged: mortality effects are directionally positive and imprecisely "
    "estimated."
)

blank()

# ─── 6.3 ─────────────────────────────────────────────────────────────────────
h2("6.3  Placebo Test (False Treatment Year 2013)")
blank()

para(
    "To assess whether the TWFE design mechanically generates spurious "
    "post-treatment effects, I assign a placebo treatment year of 2013 (two years "
    "before the true 2015 fires) and re-estimate the simple DiD within the same "
    "matched panel. Under the null, true 2013 effects should be zero because no "
    "wildfire occurred."
)
para(
    "Table 4 reports the placebo ATTs for all three thresholds. All six estimates "
    "are statistically insignificant. At T1, the placebo ATT is +0.160 for suicide "
    "(p = 0.399) and +0.140 for overdose (p = 0.488). At T3, placebo estimates are "
    "negative (-0.354 for suicide, p = 0.238; -0.351 for overdose, p = 0.202), "
    "providing particularly clean evidence against spurious divergence at the most "
    "extreme threshold."
)
para(
    "An honest reading of the T1 placebo magnitudes is warranted. The placebo ATTs "
    "(0.140-0.160 IHS units) are somewhat close in magnitude to the actual estimated "
    "ATTs (0.195-0.213 IHS units), which could be taken as evidence that the true "
    "estimates also reflect noise. The more appropriate comparison is between the "
    "actual pre-trend tests and a spurious-effect interpretation: the event-study "
    "pre-trends at T1 pass a joint F-test (p = 0.663 for suicide), and the "
    "post-treatment event-study coefficients are consistently positive across all "
    "five post-treatment years while the pre-treatment coefficients hover near zero. "
    "This temporal pattern (flat before 2015, positive after) is inconsistent "
    "with a pure noise story and consistent with a delayed causal effect. The "
    "placebo confirms that a 2013 shock generates no such temporal pattern."
)

blank()

# Table 4
caption("Table 4. Placebo Test: False Treatment Year 2013")

tbl4 = doc.add_table(rows=7, cols=4)
tbl4.style = 'Table Grid'

cell_text(tbl4, 0, 0, "", bold=True, fs=10)
cell_text(tbl4, 0, 1, "T1 (>=1,000 ac)", bold=True, center=True, fs=10)
cell_text(tbl4, 0, 2, "T2 (>=5,000 ac)", bold=True, center=True, fs=10)
cell_text(tbl4, 0, 3, "T3 (>=25,000 ac)", bold=True, center=True, fs=10)

cell_text(tbl4, 1, 0, "Panel A: IHS suicide rate", italic=True, fs=10)
for c in range(1, 4): cell_text(tbl4, 1, c, "", fs=10)

cell_text(tbl4, 2, 0, "  Placebo ATT", fs=10)
cell_text(tbl4, 2, 1, "0.160", center=True, fs=10)
cell_text(tbl4, 2, 2, "0.183", center=True, fs=10)
cell_text(tbl4, 2, 3, "-0.354", center=True, fs=10)

cell_text(tbl4, 3, 0, "  (SE)", fs=10)
cell_text(tbl4, 3, 1, "(0.189)", center=True, fs=10)
cell_text(tbl4, 3, 2, "(0.218)", center=True, fs=10)
cell_text(tbl4, 3, 3, "(0.294)", center=True, fs=10)

cell_text(tbl4, 4, 0, "Panel B: IHS overdose rate", italic=True, fs=10)
for c in range(1, 4): cell_text(tbl4, 4, c, "", fs=10)

cell_text(tbl4, 5, 0, "  Placebo ATT", fs=10)
cell_text(tbl4, 5, 1, "0.140", center=True, fs=10)
cell_text(tbl4, 5, 2, "0.222", center=True, fs=10)
cell_text(tbl4, 5, 3, "-0.351", center=True, fs=10)

cell_text(tbl4, 6, 0, "  (SE)", fs=10)
cell_text(tbl4, 6, 1, "(0.200)", center=True, fs=10)
cell_text(tbl4, 6, 2, "(0.176)", center=True, fs=10)
cell_text(tbl4, 6, 3, "(0.268)", center=True, fs=10)

note(
    "Notes: Simple DiD (Eq. 3) within each matched panel with treatment year reassigned to 2013. "
    "Estimation window: 2011-2016. SEs clustered by county. "
    "p-values: T1 suicide 0.399, T1 overdose 0.488, T2 suicide 0.404, T2 overdose 0.212, "
    "T3 suicide 0.238, T3 overdose 0.202. *** p<0.01, ** p<0.05, * p<0.10."
)

blank()

# ─── 6.4 ─────────────────────────────────────────────────────────────────────
h2("6.4  Bias-Corrected (CR2) Standard Errors")
blank()

para(
    "The primary T1 analysis clusters standard errors at the county level using the "
    "standard CR1 estimator. With 93 county clusters, CR1 is consistent, but "
    "cluster imbalance between 81 treated and 12 control counties may generate "
    "modest finite-sample distortions (Imbens and Kolesar, 2016). I re-estimate "
    "the T1 and T3 event studies replacing CR1 with CR2 (Bell-McCaffrey) "
    "bias-corrected standard errors and Satterthwaite degrees of freedom via the "
    "clubSandwich package."
)
para(
    "Table 5 presents the comparison for T1 suicide and overdose. CR2 standard "
    "errors are substantially larger than CR1 at every post-treatment horizon, "
    "reflecting the correction for cluster leverage imbalance. The largest suicide "
    "SE grows from 0.410 (CR1, k=0) to 0.822 (CR2, k=0); the largest overdose "
    "SE grows from 0.407 to 1.027. Despite these inflated SEs, all Satterthwaite "
    "p-values remain above 0.49; the mortality estimates are robustly insignificant "
    "at T1 under both CR1 and CR2 inference."
)
para(
    "For T3 (31 clusters), the Satterthwaite adjustment is severe: the effective "
    "degrees of freedom collapse to approximately one for some point estimates, "
    "making t-based inference unreliable regardless of SE size. This confirms the "
    "caution applied throughout to T3 mortality estimates. The depression "
    "cross-sectional regressions use heteroskedasticity-robust standard errors "
    "with 90 observations, for which the finite-sample correction is negligible."
)

blank()

# Table 5
caption("Table 5. CR2 Robustness: T1 Event-Study Standard Errors")

tbl5 = doc.add_table(rows=7, cols=7)
tbl5.style = 'Table Grid'

# Header
cell_text(tbl5, 0, 0, "k", bold=True, center=True, fs=10)
cell_text(tbl5, 0, 1, "Suicide\nβ̂", bold=True, center=True, fs=10)
cell_text(tbl5, 0, 2, "SE (CR1)", bold=True, center=True, fs=10)
cell_text(tbl5, 0, 3, "SE (CR2)", bold=True, center=True, fs=10)
cell_text(tbl5, 0, 4, "Overdose\nβ̂", bold=True, center=True, fs=10)
cell_text(tbl5, 0, 5, "SE (CR1)", bold=True, center=True, fs=10)
cell_text(tbl5, 0, 6, "SE (CR2)", bold=True, center=True, fs=10)

rows_data = [
    ("k=0",  "0.568", "0.410", "0.822", "0.356", "0.407", "1.027"),
    ("k=+1", "0.428", "0.409", "0.565", "0.563", "0.378", "0.564"),
    ("k=+2", "0.413", "0.415", "0.815", "0.346", "0.399", "0.968"),
    ("k=+3", "0.391", "0.411", "0.800", "0.321", "0.395", "0.978"),
    ("k=+4", "0.090", "0.099", "0.669", "0.111", "0.106", "0.837"),
]

for i, row_data in enumerate(rows_data):
    for j, val in enumerate(row_data):
        cell_text(tbl5, i+1, j, val, center=(j > 0), fs=10)

cell_text(tbl5, 6, 0, "All CR2 Satterthwaite p > 0.49", italic=True, fs=10)
for c in range(1, 7): cell_text(tbl5, 6, c, "", fs=10)

note(
    "Notes: T1 (>= 1,000 acres) event-study coefficients, reference year k = -1. "
    "CR1: standard cluster-robust; CR2: Bell-McCaffrey bias-corrected via clubSandwich package "
    "with Satterthwaite degrees of freedom. CR1 SEs from primary event-study (Eq. 2)."
)

out = r"C:\Users\chenyon\Research Paper 2026(1)\section_robustness.docx"
doc.save(out)
print("Saved: " + out)
