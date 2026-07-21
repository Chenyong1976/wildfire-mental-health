"""
Build Table 1 (Summary Statistics) and insert it into:
  - section_data.docx  (replaces the [TABLE 1 ABOUT HERE] placeholder paragraph)
  - table1.tex         (standalone LaTeX table, to be \input{} in section_data.tex)
"""
import warnings; warnings.filterwarnings('ignore')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Hard-coded stats from summary_stats_table.csv ────────────────────────────
# Format: (label, T_mean, T_sd, T_n, C_mean, C_sd, C_n, diff, pval, fmt)
PANELS = [
    {
        "title": "Panel A: County Characteristics (Pre-Treatment)",
        "rows": [
            ("WHP national rank (0–1)",
             0.863, 0.161, 81, 0.876, 0.100, 12, -0.013, 0.701, "3f"),
            ("Rural-urban continuum code (1–9)",
             4.46,  2.11,  81, 4.25,  2.77,  12,  0.21,  0.808, "1f"),
            ("Share of buildings in direct exposure zone",
             0.477, 0.201, 81, 0.402, 0.258, 12,  0.075, 0.351, "3f"),
            ("Population, 2014 (thousands)",
             46.9,  51.1,  81, 49.9,  56.6,  12, -3.0,   0.914, "1f"),
            ("Median household income, 2014 ($thousands)",
             44.5,   8.8,  81, 48.2,   7.3,  12, -3.6,   0.344, "1f"),
        ]
    },
    {
        "title": "Panel B: Pre-Treatment Outcomes (2011–2014 county means)",
        "rows": [
            ("Suicide rate (per 100,000)",
              9.2,  10.4,  81, 13.6,  10.4,  12, -4.4,  0.198, "1f"),
            ("Overdose mortality rate (per 100,000)",
             11.3,  14.8,  81, 17.7,  17.0,  12, -6.4,  0.240, "1f"),
            ("Unemployment rate (%)",
              8.9,   2.7,  79,  8.6,   3.3,  11,  0.3,  0.779, "1f"),
        ]
    },
    {
        "title": "Panel C: 2019 Outcomes",
        "rows": [
            ("Depression prevalence (%)",
             21.2,  3.6,  79, 18.0,  2.4,  11,  3.2,  0.001, "1f"),
            ("Poor mental health days (% with 14+ days/month)",
             15.1,  2.0,  79, 13.2,  1.4,  11,  1.9,  0.001, "1f"),
        ]
    },
]

def stars(p):
    if p < 0.01:  return "***"
    if p < 0.05:  return "**"
    if p < 0.10:  return "*"
    return ""

def fmt_cell(val, sd, f):
    return f"{val:.{f}}", f"({sd:.{f}})"

# ═══════════════════════════════════════════════════════════════════════════════
# 1.  LATEX TABLE
# ═══════════════════════════════════════════════════════════════════════════════
latex_lines = [
    r"\begin{table}[htbp]",
    r"\centering",
    r"\caption{Summary Statistics: T1 Matched Panel (Wildfires $\geq$1{,}000 Acres, 2015)}",
    r"\label{tab:summary_stats}",
    r"\small",
    r"\begin{tabular}{lcccccc}",
    r"\toprule",
    r" & \multicolumn{2}{c}{Treated} & \multicolumn{2}{c}{Control} & & \\",
    r"\cmidrule(lr){2-3}\cmidrule(lr){4-5}",
    r"Variable & Mean & (SD) & Mean & (SD) & Diff. & $p$-value \\",
    r"\midrule",
]

for panel in PANELS:
    latex_lines.append(r"\multicolumn{7}{l}{\textit{" +
                       panel["title"].replace("–", "--").replace("–", "--") +
                       r"}} \\")
    for (label, tm, ts, tn, cm, cs, cn, diff, pval, f) in panel["rows"]:
        mv, sv = f"{tm:.{f}}", f"{ts:.{f}}"
        mc, sc = f"{cm:.{f}}", f"{cs:.{f}}"
        d_str  = f"{diff:+.{f}}"
        p_str  = f"{pval:.3f}"
        st     = stars(pval)
        label_esc = label.replace("–", "--").replace("%", r"\%").replace("$", r"\$")
        latex_lines.append(
            f"\\quad {label_esc} & {mv} & ({sv}) & {mc} & ({sc}) & "
            f"{d_str}{st} & {p_str} \\\\"
        )

latex_lines += [
    r"\bottomrule",
    r"\end{tabular}",
    r"\begin{minipage}{\linewidth}",
    r"\vspace{4pt}",
    r"\footnotesize",
    (r"\textit{Notes:} Treated counties experienced a wildfire of at least 1{,}000 acres in 2015 "
     r"with no qualifying wildfire in the prior period (N\,=\,81). Control counties had no "
     r"qualifying wildfire in 2015--2019 (N\,=\,12). Standard deviations in parentheses. "
     r"Difference is treated minus control; $p$-values from two-sample Welch $t$-tests. "
     r"Panel~B rates are county-level means over non-suppressed CDC WONDER cells (2011--2014); "
     r"cells with fewer than ten deaths are suppressed by CDC and treated as missing. "
     r"The lower pre-treatment rates in the treated group reflect suppression-induced "
     r"selection toward higher-count cells in control counties, not pre-treatment divergence "
     r"(joint pre-trend $F$-tests: $p = 0.66$ for suicide, $p = 0.60$ for overdose). "
     r"Panel~C outcomes are from CDC PLACES 2019. "
     r"*** $p<0.01$, ** $p<0.05$, * $p<0.10$."),
    r"\end{minipage}",
    r"\end{table}",
]

latex_out = "\n".join(latex_lines)
with open(r"C:\Users\chenyon\Research Paper 2026(1)\table1.tex", "w", encoding="utf-8") as f:
    f.write(latex_out)
print("Wrote table1.tex")

# ═══════════════════════════════════════════════════════════════════════════════
# 2.  WORD TABLE  — inserted into section_data.docx
# ═══════════════════════════════════════════════════════════════════════════════
DOCX_PATH = r"C:\Users\chenyon\Research Paper 2026(1)\section_data.docx"
doc = Document(DOCX_PATH)

# ── Find the placeholder paragraph ───────────────────────────────────────────
placeholder_idx = None
for i, p in enumerate(doc.paragraphs):
    if "[TABLE 1 ABOUT HERE]" in p.text:
        placeholder_idx = i
        break
if placeholder_idx is None:
    raise ValueError("Placeholder not found in section_data.docx")

# ── Helper: set a cell's text with given font properties ─────────────────────
def set_cell(cell, text, bold=False, italic=False, center=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic

def shade_row(row, hex_color="E8E8E8"):
    """Light grey shading for panel-header rows."""
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

# ── Build the table ───────────────────────────────────────────────────────────
# 7 columns: Variable | T Mean | T (SD) | C Mean | C (SD) | Diff | p-value
COL_WIDTHS = [Inches(2.8), Inches(0.6), Inches(0.6),
              Inches(0.6), Inches(0.6), Inches(0.7), Inches(0.6)]

# Caption paragraph
cap_para = OxmlElement("w:p")
cap_run  = OxmlElement("w:r")
cap_rPr  = OxmlElement("w:rPr")
cap_b    = OxmlElement("w:b")
cap_rPr.append(cap_b)
cap_run.append(cap_rPr)
cap_t    = OxmlElement("w:t")
cap_t.text = ("Table 1. Summary Statistics: T1 Matched Panel "
              "(Wildfires ≥1,000 Acres, 2015)")
cap_run.append(cap_t)
cap_para.append(cap_run)

# Insert caption before placeholder
placeholder_elem = doc.paragraphs[placeholder_idx]._element
placeholder_elem.addprevious(cap_para)

# Create the table
tbl = doc.add_table(rows=0, cols=7)
tbl.style = "Table Grid"

# Set column widths
for i, w in enumerate(COL_WIDTHS):
    for cell in tbl.columns[i].cells:
        cell.width = w

# Header row 1: spanning headers
hdr1 = tbl.add_row()
set_cell(hdr1.cells[0], "",           bold=True, center=True)
set_cell(hdr1.cells[1], "Treated",    bold=True, center=True)
set_cell(hdr1.cells[2], "",           bold=True, center=True)
set_cell(hdr1.cells[3], "Control",    bold=True, center=True)
set_cell(hdr1.cells[4], "",           bold=True, center=True)
set_cell(hdr1.cells[5], "",           bold=True, center=True)
set_cell(hdr1.cells[6], "",           bold=True, center=True)

# Merge Treated columns
hdr1.cells[1].merge(hdr1.cells[2])
# Merge Control columns
hdr1.cells[3].merge(hdr1.cells[4])

# Header row 2: column labels
hdr2 = tbl.add_row()
set_cell(hdr2.cells[0], "Variable",   bold=True)
set_cell(hdr2.cells[1], "Mean",       bold=True, center=True)
set_cell(hdr2.cells[2], "(SD)",       bold=True, center=True)
set_cell(hdr2.cells[3], "Mean",       bold=True, center=True)
set_cell(hdr2.cells[4], "(SD)",       bold=True, center=True)
set_cell(hdr2.cells[5], "Diff.",      bold=True, center=True)
set_cell(hdr2.cells[6], "p-value",    bold=True, center=True)

# Data rows
for panel in PANELS:
    pr = tbl.add_row()
    set_cell(pr.cells[0], panel["title"], bold=True, italic=True)
    for c in pr.cells[1:]:
        set_cell(c, "")
    shade_row(pr)

    for (label, tm, ts, tn, cm, cs, cn, diff, pval, f) in panel["rows"]:
        dr = tbl.add_row()
        d_str = f"{diff:+.{f}}" + stars(pval)
        p_str = f"{pval:.3f}"
        set_cell(dr.cells[0], "  " + label)
        set_cell(dr.cells[1], f"{tm:.{f}}", center=True)
        set_cell(dr.cells[2], f"({ts:.{f}})", center=True, italic=True)
        set_cell(dr.cells[3], f"{cm:.{f}}", center=True)
        set_cell(dr.cells[4], f"({cs:.{f}})", center=True, italic=True)
        set_cell(dr.cells[5], d_str, center=True)
        set_cell(dr.cells[6], p_str, center=True)

# Obs row
obs_r = tbl.add_row()
set_cell(obs_r.cells[0], "Counties (N)", bold=True)
set_cell(obs_r.cells[1], "81", center=True, bold=True)
set_cell(obs_r.cells[2], "", center=True)
set_cell(obs_r.cells[3], "12", center=True, bold=True)
set_cell(obs_r.cells[4], "", center=True)
set_cell(obs_r.cells[5], "", center=True)
set_cell(obs_r.cells[6], "", center=True)

# ── Move table to replace placeholder ────────────────────────────────────────
# Insert table XML before the placeholder, then remove placeholder
placeholder_elem.addprevious(tbl._tbl)
placeholder_elem.getparent().remove(placeholder_elem)

# ── Notes paragraph ───────────────────────────────────────────────────────────
note_p = doc.add_paragraph()
note_p.paragraph_format.space_before = Pt(4)
note_p.paragraph_format.space_after  = Pt(0)
note_r = note_p.add_run(
    "Notes: Treated counties experienced a wildfire of at least 1,000 acres in 2015 "
    "with no qualifying wildfire in the prior period (N = 81). Control counties had no "
    "qualifying wildfire in 2015–2019 (N = 12). Standard deviations in parentheses. "
    "Difference is treated minus control; p-values from two-sample Welch t-tests. "
    "Panel B rates are county-level means over non-suppressed CDC WONDER cells (2011–2014); "
    "cells with fewer than ten deaths are suppressed and treated as missing. "
    "The lower pre-treatment mortality rates in the treated group reflect suppression-induced "
    "selection toward higher-count cells in control counties, not pre-treatment divergence "
    "(joint pre-trend F-tests: p = 0.66 for suicide, p = 0.60 for overdose). "
    "Panel C outcomes are from CDC PLACES 2019. *** p<0.01, ** p<0.05, * p<0.10."
)
note_r.font.name = "Times New Roman"
note_r.font.size = Pt(10)
note_r.italic    = True

# Move notes after the table
tbl._tbl.addnext(note_p._element)

doc.save(DOCX_PATH)
print(f"Updated: {DOCX_PATH}")
